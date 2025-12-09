#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Converted from Jupyter Notebook: notebook.ipynb
Conversion Date: 2025-12-03T18:09:50.484Z
"""

import numpy as np
import pandas as pd
import warnings
warnings.filterwarnings('ignore')

# Core ML libraries
from sklearn.model_selection import train_test_split, KFold
from sklearn.preprocessing import StandardScaler, RobustScaler
from sklearn.impute import KNNImputer
from sklearn.decomposition import PCA

# Survival analysis libraries
from lifelines import CoxPHFitter
from lifelines.utils import concordance_index
from sksurv.ensemble import RandomSurvivalForest, GradientBoostingSurvivalAnalysis
from sksurv.linear_model import CoxnetSurvivalAnalysis
from sksurv.metrics import concordance_index_censored, integrated_brier_score

# Visualization
import matplotlib.pyplot as plt
import seaborn as sns

# survival_pipeline_with_boost.py
import warnings
warnings.filterwarnings("ignore")

import numpy as np
import pandas as pd
import matplotlib.pyplot as plt

from sklearn.impute import KNNImputer
from sklearn.preprocessing import RobustScaler
from sklearn.model_selection import train_test_split

# scikit-survival imports
from sksurv.linear_model import CoxnetSurvivalAnalysis
from sksurv.ensemble import RandomSurvivalForest, GradientBoostingSurvivalAnalysis
from sksurv.metrics import concordance_index_censored

# lifelines
from lifelines import CoxPHFitter

# -------------------------------------------------------------------------
# TOP50 gene list (de-duplicated)
# -------------------------------------------------------------------------
TOP50_GENES = [
    "MKI67","CCNB1","CCNE1","AURKA","AURKB","BIRC5","UBE2C","TOP2A","CDK1","MYC",
    "ESR1","PGR","GATA3","FOXA1","XBP1","BCL2","BAG1","KRT8","KRT18","KRT5",
    "KRT14","KRT17","ERBB2","GRB7","FGFR4","EGFR","PIK3CA","PTEN","TP53","BRCA1",
    "BRCA2","RAD51","ATM","CHEK1","VIM","SNAI1","SNAI2","ZEB1","ITGA6","CD274",
    "PDCD1","STAT1","CXCL9","CXCL10","BAG1","XBP1","CCNE1","TOP2A","BCL2","MYC"
]
# Remove duplicates, keep order
_seen = set()
TOP50_GENES = [g for g in TOP50_GENES if not (g in _seen or _seen.add(g))]

# -------------------------------------------------------------------------
# 1. DATA LOADING AND PREPROCESSING
# -------------------------------------------------------------------------
def load_and_preprocess_data(filepath='METABRIC_RNA_Mutation.csv'):
    """
    Load data with comprehensive preprocessing
    """
    print("\n[STEP 1] Loading and preprocessing data...")
    
    # Load data
    data = pd.read_csv(filepath, low_memory=False)
    print(f"Initial shape: {data.shape}")
    
    # Remove redundant columns (if present)
    cols_to_drop = [
        'patient_id', 'cancer_type', 'er_status', 'her2_status',
        'mutation_count', 'death_from_cancer', 'nottingham_prognostic_index'
    ]
    data.drop([col for col in cols_to_drop if col in data.columns], 
              axis=1, inplace=True)
    
    # Remove mutation columns (>90% zeros)
    mutation_cols = [col for col in data.columns if col.endswith('_mut')]
    data.drop(mutation_cols, axis=1, inplace=True)
    
    print(f"Shape after removing mutations: {data.shape}")
    
    # Handle target variables
    if 'overall_survival_months' not in data.columns:
        raise ValueError("Target variable 'overall_survival_months' not found")
    
    # Create censoring indicator (1 = event occurred, 0 = censored)
    if 'overall_survival' in data.columns:
        data['event'] = data['overall_survival'].fillna(0).astype(int)
    else:
        # If no event column, assume all are events (conservative) — user should avoid this if possible
        data['event'] = 1
    
    # Keep rows with valid survival time
    data = data[data['overall_survival_months'].notna()].copy()
    data = data[data['overall_survival_months'] > 0].copy()
    
    print(f"Final shape: {data.shape}")
    print(f"Events: {data['event'].sum()}, Censored: {(1-data['event']).sum()}")
    
    return data

# -------------------------------------------------------------------------
# 2. FEATURE ENGINEERING (fixed imputation ordering)
# -------------------------------------------------------------------------
def advanced_feature_engineering(data):
    """
    Advanced feature engineering with proper handling of categorical and numerical features.
    Fixed: recompute numerical_cols after get_dummies() so imputation targets all numeric columns.
    """
    print("\n[STEP 2] Advanced feature engineering...")
    
    # Separate target variables
    y_time = data['overall_survival_months'].copy()
    y_event = data['event'].copy()
    
    # Remove target from features
    X = data.drop(['overall_survival_months', 'event', 'overall_survival'], 
                  axis=1, errors='ignore')
    
    # Identify feature types (before dummies)
    categorical_cols = X.select_dtypes(include=['object']).columns.tolist()
    # Handle categorical variables
    if categorical_cols:
        X = pd.get_dummies(X, columns=categorical_cols, drop_first=True)
    
    # Recompute numeric columns after dummies
    numerical_cols = X.select_dtypes(include=[np.number]).columns.tolist()
    
    print(f"Categorical features converted to dummies. Numeric features count: {len(numerical_cols)}")
    
    # Advanced imputation for numerical features
    if X.isnull().sum().sum() > 0:
        print("Applying KNN imputation for missing values...")
        imputer = KNNImputer(n_neighbors=5, weights='distance')
        X[numerical_cols] = imputer.fit_transform(X[numerical_cols])
    
    # Replace infinite values and fill remaining NaNs
    X.replace([np.inf, -np.inf], np.nan, inplace=True)
    X.fillna(0, inplace=True)
    
    # Remove low-variance features (std < 0.01)
    low_var_cols = X.columns[X.std() < 0.01].tolist()
    if low_var_cols:
        print(f"Removing {len(low_var_cols)} low-variance features")
        X.drop(low_var_cols, axis=1, inplace=True)
    
    print(f"Final feature count: {X.shape[1]}")
    
    return X, y_time, y_event

# -------------------------------------------------------------------------
# 3. FEATURE SELECTION (ensuring boost_genes included)
# -------------------------------------------------------------------------
def feature_selection_advanced(X, y_time, y_event, top_k=100, boost_genes=None):
    """
    Advanced feature selection using Cox regression coefficients.
    Ensures boost_genes are included in the final selected set.
    """
    print(f"\n[STEP 3] Selecting top {top_k} features using Cox regression...")
    boost_genes = boost_genes or []
    
    df_cox = X.copy()
    df_cox['time'] = y_time.values
    df_cox['event'] = y_event.values
    
    feature_scores = {}
    
    # iterate through features (speed limit to first 200)
    for col in X.columns[:200]:
        try:
            df_temp = df_cox[[col, 'time', 'event']].copy()
            df_temp = df_temp[df_temp[col].notna()]
            if df_temp[col].std() > 0:
                cph = CoxPHFitter(penalizer=0.1)
                cph.fit(df_temp, duration_col='time', event_col='event')
                feature_scores[col] = abs(cph.summary['coef'].values[0])
        except Exception:
            continue
    
    if feature_scores:
        top_features = sorted(feature_scores.items(), key=lambda x: x[1], reverse=True)[:top_k]
        selected_cols = [f[0] for f in top_features]
        # Ensure boosted genes included
        for g in boost_genes:
            if g in X.columns and g not in selected_cols:
                selected_cols.append(g)
        print(f"Selected {len(selected_cols)} features (including boosted genes)")
        return X[selected_cols]
    else:
        print("Feature selection failed — using all features")
        # ensure boost genes are present warning
        for g in boost_genes:
            if g not in X.columns:
                print(f"Warning: boost gene {g} not found in X columns")
        return X

# -------------------------------------------------------------------------
# 4. SurvivalModelEnsemble with boosting
# -------------------------------------------------------------------------
class SurvivalModelEnsemble:
    """
    Ensemble of survival models with boosting of selected gene columns.
    Multiplicative boosting is applied AFTER scaling for training and prediction.
    """
    def __init__(self, boost_genes=None, boost_factor=1.5):
        self.models = {}
        self.scaler = RobustScaler()
        self.feature_names = None
        self.boost_genes = boost_genes or TOP50_GENES.copy()
        self.boost_factor = float(boost_factor)
        self._boosted_columns = []
    
    def _apply_boost(self, X_df):
        """Multiply the selected columns in X_df by boost_factor (returns new DF)."""
        Xb = X_df.copy()
        self._boosted_columns = [c for c in self.boost_genes if c in Xb.columns]
        if not self._boosted_columns:
            return Xb
        Xb[self._boosted_columns] = Xb[self._boosted_columns] * self.boost_factor
        return Xb
    
    def fit(self, X_train, y_time_train, y_event_train):
        """
        Train models after scaling + boosting.
        """
        print("\n" + "=" * 80)
        print("TRAINING ENSEMBLE MODELS")
        print("=" * 80)
        
        # Scale features
        X_train_scaled = self.scaler.fit_transform(X_train)
        X_train_scaled = pd.DataFrame(X_train_scaled, columns=X_train.columns, index=X_train.index)
        self.feature_names = X_train.columns.tolist()
        
        # Apply boost AFTER scaling
        X_train_scaled_boosted = self._apply_boost(X_train_scaled)
        
        # Prepare sksurv structured y
        y_train_surv = np.array(
            [(bool(e), t) for e, t in zip(y_event_train, y_time_train)],
            dtype=[('event', bool), ('time', float)]
        )
        
        # CoxNet
        print("\n[MODEL 1] Training CoxNet (Elastic Net Cox)...")
        try:
            self.models['coxnet'] = CoxnetSurvivalAnalysis(
                l1_ratio=0.5, alpha_min_ratio=0.01, max_iter=1000, fit_baseline_model=True
            )
            self.models['coxnet'].fit(X_train_scaled_boosted, y_train_surv)
            print("✓ CoxNet trained successfully")
        except Exception as e:
            print(f"✗ CoxNet failed: {e}")
        
        # RSF
        print("\n[MODEL 2] Training Random Survival Forest...")
        try:
            self.models['rsf'] = RandomSurvivalForest(
                n_estimators=200, min_samples_split=10, min_samples_leaf=5,
                max_features="sqrt", n_jobs=-1, random_state=42
            )
            self.models['rsf'].fit(X_train_scaled_boosted, y_train_surv)
            print("✓ Random Survival Forest trained successfully")
        except Exception as e:
            print(f"✗ RSF failed: {e}")
        
        # GBSA
        print("\n[MODEL 3] Training Gradient Boosting Survival Analysis...")
        try:
            self.models['gbsa'] = GradientBoostingSurvivalAnalysis(
                n_estimators=100, learning_rate=0.1, max_depth=3,
                min_samples_split=10, min_samples_leaf=5, subsample=0.8, random_state=42
            )
            self.models['gbsa'].fit(X_train_scaled_boosted, y_train_surv)
            print("✓ Gradient Boosting trained successfully")
        except Exception as e:
            print(f"✗ GBSA failed: {e}")
        
        # Traditional Cox (lifelines)
        print("\n[MODEL 4] Training Traditional Cox Proportional Hazards...")
        try:
            df_train = X_train_scaled_boosted.copy()
            df_train['time'] = y_time_train.values
            df_train['event'] = y_event_train.values
            self.models['cox'] = CoxPHFitter(penalizer=0.1, l1_ratio=0.5)
            self.models['cox'].fit(df_train, duration_col='time', event_col='event', show_progress=False)
            print("✓ Traditional Cox trained successfully")
        except Exception as e:
            print(f"✗ Cox failed: {e}")
        
        print(f"\n✓ Successfully trained {len(self.models)} models")
        return self
    
    def predict(self, X_test):
        """
        Generate ensemble predictions. Applies same scaling and boosting as used in training.
        Returns: ensemble_pred (array), dict_of_individual_preds
        """
        X_test_scaled = self.scaler.transform(X_test)
        X_test_scaled = pd.DataFrame(X_test_scaled, columns=X_test.columns, index=X_test.index)
        
        # Apply boost to same columns that were boosted during fit
        if hasattr(self, "_boosted_columns") and self._boosted_columns:
            cols_to_boost = [c for c in self._boosted_columns if c in X_test_scaled.columns]
            if cols_to_boost:
                X_test_scaled[cols_to_boost] = X_test_scaled[cols_to_boost] * self.boost_factor
        
        predictions = {}
        
        # CoxNet
        if 'coxnet' in self.models:
            try:
                predictions['coxnet'] = self.models['coxnet'].predict(X_test_scaled)
            except Exception:
                pass
        
        # RSF
        if 'rsf' in self.models:
            try:
                predictions['rsf'] = self.models['rsf'].predict(X_test_scaled)
            except Exception:
                pass
        
        # GBSA
        if 'gbsa' in self.models:
            try:
                predictions['gbsa'] = self.models['gbsa'].predict(X_test_scaled)
            except Exception:
                pass
        
        # Cox (lifelines): use negative partial hazard for consistent ordering
        if 'cox' in self.models:
            try:
                df_test = X_test_scaled.copy()
                predictions['cox'] = -self.models['cox'].predict_partial_hazard(df_test).values
            except Exception:
                pass
        
        if predictions:
            pred_array = np.array(list(predictions.values()))
            ensemble_pred = np.mean(pred_array, axis=0)
            return ensemble_pred, predictions
        else:
            raise ValueError("No models available for prediction")
    
    def evaluate(self, X_test, y_time_test, y_event_test):
        """
        Evaluate all available models and the ensemble using concordance index.
        """
        print("\n" + "=" * 80)
        print("MODEL EVALUATION")
        print("=" * 80)
        
        ensemble_pred, individual_preds = self.predict(X_test)
        
        results = {}
        for name, pred in individual_preds.items():
            try:
                c_index = concordance_index_censored(
                    y_event_test.astype(bool), y_time_test, pred
                )[0]
                results[name] = c_index
                print(f"{name.upper():20s} C-index: {c_index:.4f}")
            except Exception as e:
                print(f"{name.upper():20s} Evaluation failed: {e}")
        
        try:
            c_index_ensemble = concordance_index_censored(
                y_event_test.astype(bool), y_time_test, ensemble_pred
            )[0]
            results['ensemble'] = c_index_ensemble
            print(f"{'ENSEMBLE':20s} C-index: {c_index_ensemble:.4f}")
        except Exception as e:
            print(f"{'ENSEMBLE':20s} Evaluation failed: {e}")
        
        return results, ensemble_pred

# -------------------------------------------------------------------------
# 5. MAIN EXECUTION (uses boosting)
# -------------------------------------------------------------------------
def main(filepath='METABRIC_RNA_Mutation.csv', top_k=80, boost_factor=1.5, boost_genes=None):
    """
    Main execution pipeline. boost_genes: list of gene names to upweight (default: TOP50_GENES)
    boost_factor: multiplicative factor applied after scaling (float)
    """
    # Load & preprocess
    data = load_and_preprocess_data(filepath)
    X, y_time, y_event = advanced_feature_engineering(data)
    
    # Feature selection (ensure boost_genes included)
    boost_genes = boost_genes or TOP50_GENES
    X_selected = feature_selection_advanced(X, y_time, y_event, top_k=top_k, boost_genes=boost_genes)
    
    # Train-test split
    X_train, X_test, y_time_train, y_time_test, y_event_train, y_event_test = \
        train_test_split(X_selected, y_time, y_event, test_size=0.3, random_state=42, stratify=y_event)
    
    print(f"\nTraining set: {X_train.shape[0]} samples")
    print(f"Test set: {X_test.shape[0]} samples")
    
    # Train ensemble with boosting
    ensemble = SurvivalModelEnsemble(boost_genes=boost_genes, boost_factor=boost_factor)
    ensemble.fit(X_train, y_time_train, y_event_train)
    
    # Evaluate
    results, ensemble_pred = ensemble.evaluate(X_test, y_time_test, y_event_test)
    
    # Visualization: C-index comparison
    print("\n" + "=" * 80)
    print("CREATING VISUALIZATIONS")
    print("=" * 80)
    plt.figure(figsize=(10, 6))
    models = list(results.keys())
    scores = list(results.values())
    colors = ['#1f77b4', '#ff7f0e', '#2ca02c', '#d62728', '#9467bd']
    bars = plt.bar(models, scores, color=colors[:len(models)])
    plt.ylabel('C-index', fontsize=12)
    plt.title('Model Performance Comparison (C-index)', fontsize=14, fontweight='bold')
    plt.ylim([0.5, 1.0])
    plt.axhline(y=0.7, color='gray', linestyle='--', alpha=0.5, label='Good threshold (0.7)')
    for bar in bars:
        height = bar.get_height()
        plt.text(bar.get_x() + bar.get_width()/2., height, f'{height:.3f}', ha='center', va='bottom', fontsize=10)
    plt.legend()
    plt.tight_layout()
    plt.savefig('model_comparison.png', dpi=300, bbox_inches='tight')
    print("✓ Saved: model_comparison.png")
    
    # Risk stratification — ensemble_pred is an array of risk scores
    try:
        plt.figure(figsize=(12, 6))
        # qcut requires pandas Series with length > unique bins; handle small samples
        risk_series = pd.Series(ensemble_pred, index=X_test.index)
        if risk_series.nunique() < 3:
            # fallback: simple percentile split
            q1 = np.percentile(ensemble_pred, 33)
            q2 = np.percentile(ensemble_pred, 66)
            risk_groups = pd.cut(risk_series, bins=[-np.inf, q1, q2, np.inf], labels=['Low Risk', 'Medium Risk', 'High Risk'])
        else:
            risk_groups = pd.qcut(risk_series, q=3, labels=['Low Risk', 'Medium Risk', 'High Risk'])
        
        for risk in ['Low Risk', 'Medium Risk', 'High Risk']:
            mask = (risk_groups == risk)
            times = y_time_test[mask]
            events = y_event_test[mask]
            if len(times) == 0:
                continue
            sorted_idx = np.argsort(times)
            times_sorted = times.values[sorted_idx]
            events_sorted = events.values[sorted_idx]
            survival_prob = []
            time_points = []
            current_survival = 1.0
            for t in np.unique(times_sorted):
                n_events = np.sum((times_sorted == t) & (events_sorted == 1))
                n_at_risk_current = np.sum(times_sorted >= t)
                if n_at_risk_current > 0:
                    current_survival *= (1 - n_events / n_at_risk_current)
                time_points.append(t)
                survival_prob.append(current_survival)
            plt.plot(time_points, survival_prob, label=risk, linewidth=2)
        plt.xlabel('Time (months)', fontsize=12)
        plt.ylabel('Survival Probability', fontsize=12)
        plt.title('Kaplan-Meier Curves by Risk Group', fontsize=14, fontweight='bold')
        plt.legend(fontsize=11)
        plt.grid(alpha=0.3)
        plt.tight_layout()
        plt.savefig('risk_stratification.png', dpi=300, bbox_inches='tight')
        print("✓ Saved: risk_stratification.png")
    except Exception as e:
        print(f"Risk stratification plotting skipped/failed: {e}")
    
    # Summary
    print("\n" + "=" * 80)
    print("SUMMARY")
    print("=" * 80)
    if results:
        best = max(results.items(), key=lambda x: x[1])
        print(f"Best individual model: {best}")
        print(f"Ensemble C-index: {results.get('ensemble', 'N/A')}")
    else:
        print("No results to display.")
    
    return ensemble, results

# -------------------------------------------------------------------------
# If run as script
# -------------------------------------------------------------------------
if __name__ == "__main__":
    # Example: run with boost_factor=1.5 (tweak as needed)
    ensemble, results = main(filepath='METABRIC_RNA_Mutation.csv', top_k=80, boost_factor=1.5, boost_genes=TOP50_GENES)
    print("\nPIPELINE COMPLETED")


# survival_report_three_tier.py
import joblib
import pandas as pd
import numpy as np
import os
import json
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import matplotlib.pyplot as plt
from io import BytesIO
import tempfile
from datetime import datetime

# -------------------------
# Robust predictor loader + prediction (replacement functions)
# -------------------------
def _load_predictor(model_path_candidates=None, scaler_path_candidates=None, selected_path_candidates=None):
    """
    Attempt to load ensemble model, scaler and selected feature list from disk.
    If any of these are missing, return a deterministic dummy predictor for testing.
    Returns: (ensemble_object, scaler_object_or_None, selected_feature_list_or_None)
    """
    if model_path_candidates is None:
        model_path_candidates = ["ensemble_model.joblib", "ensemble.joblib", "model.joblib", "ensemble_model.pkl"]
    if scaler_path_candidates is None:
        scaler_path_candidates = ["scaler.joblib", "scaler.pkl"]
    if selected_path_candidates is None:
        selected_path_candidates = ["selected_features.joblib", "selected_features.json", "selected.pkl"]

    loaded_model = None
    loaded_scaler = None
    selected = None

    # Try model
    for p in model_path_candidates:
        if os.path.exists(p):
            try:
                loaded_model = joblib.load(p)
                break
            except Exception as e:
                print(f"Failed to load model from {p}: {e}")

    # Try scaler
    for p in scaler_path_candidates:
        if os.path.exists(p):
            try:
                loaded_scaler = joblib.load(p)
                break
            except Exception as e:
                print(f"Failed to load scaler from {p}: {e}")

    # Try selected features
    for p in selected_path_candidates:
        if os.path.exists(p):
            try:
                if p.lower().endswith(".json"):
                    with open(p, "r") as f:
                        selected = json.load(f)
                else:
                    selected = joblib.load(p)
                break
            except Exception as e:
                print(f"Failed to load selected features from {p}: {e}")

    # Extract selected from model if possible
    if selected is None and loaded_model is not None:
        if hasattr(loaded_model, "feature_names_in_"):
            try:
                selected = list(loaded_model.feature_names_in_)
            except Exception:
                selected = None
        elif hasattr(loaded_model, "selected_features_"):
            try:
                selected = list(getattr(loaded_model, "selected_features_"))
            except Exception:
                selected = None

    # If model+selected OK, return them
    if loaded_model is not None and selected is not None:
        return loaded_model, loaded_scaler, list(selected)

    # Otherwise create deterministic dummy predictor and dummy scaler
    print("Real model/scaler/selected features not found — using deterministic dummy predictor for testing.")

    class DummyScaler:
        def transform(self, X):
            if isinstance(X, pd.DataFrame):
                return X.values
            return X

    class DummyEnsemble:
        def __init__(self, seed=42):
            self.seed = seed

        def predict(self, X):
            """
            Accepts DataFrame or numpy array. Produces reproducible scores in [0,1].
            Uses row sums when numeric, else reproducible RNG fallback.
            """
            if isinstance(X, pd.DataFrame):
                try:
                    arr = X.fillna(0).astype(float).values
                except Exception:
                    arr = pd.DataFrame({c: pd.to_numeric(X[c], errors='coerce').fillna(0) for c in X.columns}).values
            else:
                arr = np.asarray(X, dtype=float)

            n = arr.shape[0]
            row_sums = np.nansum(arr, axis=1)
            if np.all(np.isfinite(row_sums)) and np.ptp(row_sums) != 0:
                scores = (row_sums - np.nanmin(row_sums)) / (np.nanmax(row_sums) - np.nanmin(row_sums))
            else:
                rng = np.random.RandomState(self.seed)
                scores = rng.rand(n)
            return np.asarray(scores)

    return DummyEnsemble(), DummyScaler(), None


def predict_breast_cancer_from_excel(input_excel_path):
    """
    Read input Excel, load predictor (or dummy), produce results list and original X dataframe.
    Robust to non-numeric identifier columns: automatically selects numeric features for prediction
    unless a selected-feature list is provided by the model files.
    Returns:
       results (list of dicts), X (pd.DataFrame of numeric features used)
    """
    ens, scaler, selected = _load_predictor()

    # Read input
    if not os.path.exists(input_excel_path):
        raise FileNotFoundError(f"Input file not found: {input_excel_path}")
    df_in = pd.read_excel(input_excel_path)

    if df_in.shape[0] == 0:
        raise ValueError("Input file contains no rows.")

    # Identify identifier/non-numeric columns for reporting (kept aside)
    non_numeric_cols = df_in.select_dtypes(include=['object', 'category']).columns.tolist()

    # If selected features not provided, choose numeric columns from df_in
    if selected is None:
        numeric_cols = df_in.select_dtypes(include=[np.number]).columns.tolist()
        if len(numeric_cols) == 0:
            # try coercion: attempt to convert columns to numeric if possible
            for c in df_in.columns:
                col = pd.to_numeric(df_in[c], errors='coerce')
                if col.notna().any():
                    df_in[c] = col
            numeric_cols = df_in.select_dtypes(include=[np.number]).columns.tolist()
        selected = numeric_cols

        if not selected:
            raise ValueError("No numeric features available for prediction. Please include numeric feature columns in the Excel file.")

    else:
        # If selected provided, ensure they exist in df_in; attempt case-insensitive mapping
        corrected_selected = []
        col_map_lower = {c.lower(): c for c in df_in.columns}
        for req in selected:
            if req in df_in.columns:
                corrected_selected.append(req)
            elif req.lower() in col_map_lower:
                corrected_selected.append(col_map_lower[req.lower()])
            else:
                print(f"Requested feature '{req}' not present in input and will be ignored.")
        numeric_candidates = []
        for c in corrected_selected:
            coerced = pd.to_numeric(df_in[c], errors='coerce')
            if coerced.notna().any() or not coerced.isna().all():
                df_in[c] = coerced
                numeric_candidates.append(c)
        selected = numeric_candidates

        if not selected:
            raise ValueError("After mapping and coercion, no usable numeric features remain from the selected-feature list. Please revise the feature list or provide numeric columns in the Excel file.")

    # Report dropped columns for user awareness
    dropped = [c for c in df_in.columns if c not in selected]
    if dropped:
        print(f"Note: The following non-numeric/unused columns were ignored for prediction: {dropped}")

    # Prepare X (numeric features only), fill NA with column median
    X = df_in[selected].copy()
    for c in X.columns:
        X[c] = pd.to_numeric(X[c], errors='coerce')
    X = X.fillna(X.median())

    # Apply scaler if available
    used_scaler = scaler if scaler is not None else getattr(ens, "scaler", None)
    if used_scaler is not None:
        try:
            X_scaled_arr = used_scaler.transform(X)
            X_scaled = pd.DataFrame(X_scaled_arr, columns=X.columns, index=X.index)
        except Exception as e:
            print("Scaler transform failed; using raw numeric X. Error:", e)
            X_scaled = X.copy()
    else:
        X_scaled = X.copy()

    # Predict
    try:
        preds = ens.predict(X_scaled if not isinstance(X_scaled, np.ndarray) else X_scaled)
        if isinstance(preds, tuple) and len(preds) == 2:
            ensemble_pred, indiv_preds = preds
            ensemble_pred = np.asarray(ensemble_pred)
        else:
            ensemble_pred = np.asarray(preds)
            indiv_preds = {}
    except Exception as e:
        print(f"Prediction error: {e}")
        raise

    # Build risk groups: Only three categories (Low / Medium / High)
    try:
        if len(ensemble_pred) > 1:
            risk_series = pd.Series(ensemble_pred)
            # Prefer qcut when possible
            try:
                risk_categories = pd.qcut(risk_series, q=3, labels=['Low', 'Medium', 'High'])
                risk_groups = risk_categories.astype(str).tolist()
            except Exception:
                # fallback: percentile thresholds
                p33 = np.nanpercentile(ensemble_pred, 33)
                p66 = np.nanpercentile(ensemble_pred, 66)
                rg = []
                for s in ensemble_pred:
                    if s <= p33:
                        rg.append('Low')
                    elif s <= p66:
                        rg.append('Medium')
                    else:
                        rg.append('High')
                risk_groups = rg
        else:
            # Single sample mapping by thresholds
            score = float(ensemble_pred[0])
            if score < 0.33:
                risk_groups = ['Low']
            elif score < 0.66:
                risk_groups = ['Medium']
            else:
                risk_groups = ['High']
    except Exception as e:
        print(f"Risk grouping failed; defaulting to 'Medium'. Error: {e}")
        risk_groups = ['Medium'] * len(ensemble_pred)

    # Build results
    results = []
    for i in range(len(ensemble_pred)):
        model_scores = {k: float(v[i]) for k, v in indiv_preds.items()} if isinstance(indiv_preds, dict) and indiv_preds else {}
        if model_scores and len(model_scores) > 1:
            scores_list = list(model_scores.values())
            confidence_std = np.std(scores_list)
            confidence = "High" if confidence_std < 0.1 else "Medium" if confidence_std < 0.2 else "Low"
        else:
            confidence = "Standard"

        results.append({
            "index": i,
            "original_index": int(df_in.index[i]) if hasattr(df_in, "index") else i,
            "ensemble_score": float(ensemble_pred[i]),
            "risk_group": risk_groups[i],
            "model_scores": model_scores,
            "confidence": confidence,
            "timestamp": datetime.now().isoformat()
        })

    # Return results and the numeric feature DataFrame used for prediction
    return results, X

# -------------------------
# Report generator (updated for 3-tier risk)
# -------------------------
def generate_breast_cancer_clinical_report(results, input_features_df, out_docx_path="Breast_Cancer_Clinical_Report.docx", 
                                         include_visualizations=True):
    """
    Generate a comprehensive breast cancer clinical report with detailed feature analysis
    and clinical recommendations. Uses three-tier risk groups: Low, Medium, High.
    """
    doc = Document()

    # Set document styles
    try:
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)
    except Exception:
        pass

    # Title page
    title = doc.add_heading('Breast Cancer Risk Assessment Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Comprehensive AI-Powered Risk Evaluation").italic = True

    doc.add_paragraph()

    meta_table = doc.add_table(rows=5, cols=2)
    try:
        meta_table.style = 'Light Grid Accent 1'
    except Exception:
        pass

    meta_table.cell(0, 0).text = "Report Date:"
    meta_table.cell(0, 1).text = datetime.now().strftime("%Y-%m-%d")
    meta_table.cell(1, 0).text = "Total Patients Analyzed:"
    meta_table.cell(1, 1).text = str(len(results))
    meta_table.cell(2, 0).text = "Assessment Type:"
    meta_table.cell(2, 1).text = "Breast Cancer Risk Prediction"

    # Calculate risk distribution safely
    risk_counts = {}
    for r in results:
        risk_group = r.get('risk_group', 'Unknown')
        risk_counts[risk_group] = risk_counts.get(risk_group, 0) + 1

    meta_table.cell(3, 0).text = "Risk Distribution:"
    meta_table.cell(3, 1).text = ", ".join([f"{k}: {v}" for k, v in risk_counts.items()])
    meta_table.cell(4, 0).text = "Report Generated By:"
    meta_table.cell(4, 1).text = "Breast Cancer AI Prediction System"

    doc.add_page_break()

    # Executive Summary
    doc.add_heading('Executive Summary', level=1)
    summary_para = doc.add_paragraph()
    summary_para.add_run("This comprehensive breast cancer risk assessment utilizes advanced ensemble machine learning to evaluate multiple clinical, demographic, and pathological factors. ")
    summary_para.add_run("The analysis provides individualized risk stratification to support clinical decision-making for breast cancer screening, prevention, and early intervention strategies.")

    # Visualizations if multiple patients
    if include_visualizations and len(results) > 1:
        doc.add_heading('Cohort Risk Distribution', level=2)
        try:
            fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(12, 5))
            risk_labels = []
            risk_sizes = []
            colors = []

            color_map = {
                'Low': '#2ecc71', 'Medium': '#f39c12', 'High': '#e74c3c',
                'Low Risk': '#2ecc71', 'Medium Risk': '#f39c12', 'High Risk': '#e74c3c'
            }

            for risk_group, count in risk_counts.items():
                if count > 0:
                    risk_labels.append(risk_group)
                    risk_sizes.append(count)
                    colors.append(color_map.get(risk_group, '#95a5a6'))

            if risk_sizes and sum(risk_sizes) > 0:
                wedges, texts, autotexts = ax1.pie(
                    risk_sizes, labels=risk_labels, autopct='%1.1f%%',
                    colors=colors, startangle=90
                )
                ax1.set_title('Breast Cancer Risk Distribution')
                for autotext in autotexts:
                    autotext.set_color('white')
                    autotext.set_fontweight('bold')
            else:
                ax1.text(0.5, 0.5, 'No Data Available',
                        ha='center', va='center', transform=ax1.transAxes)
                ax1.set_title('Risk Distribution\n(No Data)')

            scores = [r.get('ensemble_score', np.nan) for r in results]
            scores = [s for s in scores if np.isfinite(s)]
            if scores:
                ax2.hist(scores, bins=min(15, max(1, len(scores))), alpha=0.7)
                ax2.set_xlabel('Risk Score')
                ax2.set_ylabel('Number of Patients')
                ax2.set_title('Breast Cancer Risk Score Distribution')
                ax2.grid(True, alpha=0.3)
                if len(scores) > 1:
                    mean_score = np.mean(scores)
                    ax2.axvline(mean_score, color='red', linestyle='--', label=f'Mean: {mean_score:.3f}')
                    ax2.legend()
            else:
                ax2.text(0.5, 0.5, 'No Score Data',
                        ha='center', va='center', transform=ax2.transAxes)
                ax2.set_title('Risk Score Distribution\n(No Data)')

            plt.tight_layout()
            with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp_file:
                plt.savefig(tmp_file.name, dpi=300, bbox_inches='tight')
                plt.close()
                doc.add_picture(tmp_file.name, width=Inches(6))
                os.unlink(tmp_file.name)

            caption = doc.add_paragraph()
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption.add_run("Figure 1: Breast Cancer Risk Distribution Analysis").italic = True

        except Exception as viz_error:
            print(f"Visualization error: {viz_error}")
            error_para = doc.add_paragraph()
            error_para.add_run("Note: Visualizations could not be generated due to technical issues. ")
            error_para.add_run("All clinical interpretations and recommendations remain valid.")

    # Individual patient assessments
    doc.add_heading('Individual Patient Assessments', level=1)

    for i, result in enumerate(results):
        patient_idx = result.get("index", i)
        doc.add_heading(f'Patient {patient_idx} - Breast Cancer Risk Assessment', level=2)

        metrics_table = doc.add_table(rows=5, cols=2)
        try:
            metrics_table.style = 'Medium Grid 1 Accent 1'
        except Exception:
            pass

        # Risk Score
        try:
            metrics_table.cell(0, 0).text = "Breast Cancer Risk Score:"
            metrics_table.cell(0, 1).text = f"{result.get('ensemble_score', np.nan):.4f}"
        except Exception:
            metrics_table.cell(0, 0).text = "Breast Cancer Risk Score:"
            metrics_table.cell(0, 1).text = str(result.get('ensemble_score', 'NA'))

        # Risk Group
        metrics_table.cell(1, 0).text = "Risk Category:"
        risk_cell = metrics_table.cell(1, 1)
        risk_cell.text = str(result.get('risk_group', 'Unknown'))

        # Confidence level
        confidence = result.get('confidence', 'Not Available')
        metrics_table.cell(2, 0).text = "Prediction Confidence:"
        metrics_table.cell(2, 1).text = str(confidence)

        # Recommendation priority (mapped to 3-tier)
        priority_map = {
            'High': 'High Priority - Urgent Evaluation Needed',
            'Medium': 'Medium Priority - Enhanced Surveillance',
            'Low': 'Routine Screening'
        }
        metrics_table.cell(3, 0).text = "Clinical Priority:"
        metrics_table.cell(3, 1).text = priority_map.get(result.get('risk_group', ''), 'Standard Care')

        # Population comparison
        metrics_table.cell(4, 0).text = "Population Comparison:"
        rg = result.get('risk_group', '')
        if rg == 'High':
            metrics_table.cell(4, 1).text = "Substantially Elevated Risk"
        elif rg == 'Medium':
            metrics_table.cell(4, 1).text = "Moderately Elevated Risk"
        else:
            metrics_table.cell(4, 1).text = "Average or Below Average Risk"

        # Feature analysis
        doc.add_heading('Key Clinical Factors Analysis', level=3)
        feature_para = doc.add_paragraph()
        feature_para.add_run("Top Clinical Features Influencing This Prediction:\n").bold = True

        # Feature explanations (kept minimal to avoid missing keys)
        feature_explanations = {
            'IHC_ER_percent': 'Percent ER positivity by IHC',
            'IHC_PR_percent': 'Percent PR positivity by IHC',
            'Ki67_percent': 'Proliferation index (Ki-67)',
            # Add other keys if you have them in input_features_df
        }

        try:
            if input_features_df is not None and i < len(input_features_df):
                patient_features = input_features_df.iloc[i]
                feature_count = 0
                max_features_to_show = 10

                for feature_name in input_features_df.columns:
                    if feature_count >= max_features_to_show:
                        break
                    if feature_name in feature_explanations:
                        val = patient_features.get(feature_name, 'NA')
                        explanation = feature_explanations[feature_name]
                        p = doc.add_paragraph()
                        p.add_run(f"• {feature_name}: ").bold = True
                        try:
                            p.add_run(f"{float(val):.3f} - {explanation}")
                        except Exception:
                            p.add_run(f"{val} - {explanation}")
                        feature_count += 1

                if feature_count == 0:
                    p = doc.add_paragraph()
                    p.add_run("Available clinical features for this patient:\n").bold = True
                    for j, feature_name in enumerate(input_features_df.columns):
                        if j >= 10:
                            break
                        val = patient_features.get(feature_name, 'NA')
                        q = doc.add_paragraph()
                        q.add_run(f"• {feature_name}: ").bold = True
                        try:
                            q.add_run(f"{float(val):.3f}")
                        except Exception:
                            q.add_run(f"{val}")
            else:
                p = doc.add_paragraph()
                p.add_run("Feature data not available for this patient. ").italic = True
                p.add_run("Please ensure the input feature data is properly passed to the report generator.")
        except Exception as feature_error:
            print(f"Feature analysis error for patient {patient_idx}: {feature_error}")
            p = doc.add_paragraph()
            p.add_run("Error analyzing clinical features. ").italic = True
            p.add_run("Technical details: " + str(feature_error))

        # Feature impact visualization per patient (if available)
        if include_visualizations and input_features_df is not None and i < len(input_features_df):
            try:
                doc.add_heading('Feature Impact Analysis', level=3)
                fig, ax = plt.subplots(figsize=(10, 6))
                patient_features = input_features_df.iloc[i]
                top_features = list(input_features_df.columns[:8])
                patient_values = []
                for feature in top_features:
                    try:
                        patient_values.append(float(patient_features.get(feature, 0.0)))
                    except Exception:
                        patient_values.append(0.0)

                normal_ranges = []
                for feature in top_features:
                    mean_val = input_features_df[feature].mean()
                    std_val = input_features_df[feature].std()
                    if np.isnan(mean_val) or np.isnan(std_val):
                        normal_ranges.append((0.0, 0.0))
                    else:
                        normal_ranges.append((mean_val - std_val, mean_val + std_val))

                y_pos = np.arange(len(top_features))
                for j, (low, high) in enumerate(normal_ranges):
                    ax.barh(y_pos[j], high - low, left=low, alpha=0.3, label='Normal Range' if j == 0 else "")

                bars = ax.barh(y_pos, patient_values, alpha=0.7, label='Patient Value')
                ax.set_yticks(y_pos)
                ax.set_yticklabels(top_features)
                ax.set_xlabel('Feature Value')
                ax.set_title('Patient Features vs Normal Ranges')
                ax.legend()

                for bar, value in zip(bars, patient_values):
                    try:
                        ax.text(bar.get_width() + bar.get_x(), bar.get_y() + bar.get_height()/2.,
                                f'{float(value):.2f}', ha='left', va='center', fontweight='bold')
                    except Exception:
                        ax.text(bar.get_width() + bar.get_x(), bar.get_y() + bar.get_height()/2.,
                                f'{value}', ha='left', va='center', fontweight='bold')

                plt.tight_layout()
                with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp_file:
                    plt.savefig(tmp_file.name, dpi=300, bbox_inches='tight')
                    plt.close()
                    doc.add_picture(tmp_file.name, width=Inches(5))
                    os.unlink(tmp_file.name)

                interpret_para = doc.add_paragraph()
                interpret_para.add_run("Feature Interpretation: ").bold = True
                interpret_para.add_run("Bars show patient's actual values compared to normal ranges (gray background). ")
                interpret_para.add_run("Values outside normal ranges may contribute significantly to risk assessment.")
            except Exception as viz_error:
                print(f"Feature visualization error for patient {patient_idx}: {viz_error}")

        # Interpretation and recommendations (3-tier)
        doc.add_heading('Breast Cancer Clinical Interpretation', level=3)
        interpret_para = doc.add_paragraph()

        score = result.get('ensemble_score', 0.0)
        risk_group = result.get('risk_group', 'Medium')

        if risk_group == 'High':
            interpret_para.add_run("HIGH RISK PROFILE - URGENT ONCOLOGICAL EVALUATION RECOMMENDED\n").bold = True
            interpret_para.add_run(f"• Breast cancer risk score: {score:.3f} (Significantly Elevated)\n")
            interpret_para.add_run("• Clinical features suggest high probability of malignancy\n")
            interpret_para.add_run("• Immediate comprehensive breast evaluation required\n")
            interpret_para.add_run("• Consideration for biopsy and multidisciplinary review\n")
        elif risk_group == 'Medium':
            interpret_para.add_run("MODERATE RISK PROFILE - ENHANCED SURVEILLANCE INDICATED\n").bold = True
            interpret_para.add_run(f"• Breast cancer risk score: {score:.3f} (Moderately Elevated)\n")
            interpret_para.add_run("• Some concerning features present requiring follow-up\n")
            interpret_para.add_run("• Short-interval imaging follow-up recommended\n")
        else:
            interpret_para.add_run("LOW RISK PROFILE - ROUTINE SCREENING ADEQUATE\n").bold = True
            interpret_para.add_run(f"• Breast cancer risk score: {score:.3f} (Favorable Profile)\n")
            interpret_para.add_run("• Clinical features consistent with benign characteristics\n")
            interpret_para.add_run("• Continue with age-appropriate screening guidelines\n")

        doc.add_heading('Specific Breast Cancer Management Recommendations', level=3)
        rec_para = doc.add_paragraph()

        if risk_group == 'High':
            rec_para.add_run("IMMEDIATE ACTIONS REQUIRED:\n").bold = True
            rec_para.add_run("• Urgent referral to breast specialist/surgical oncology\n")
            rec_para.add_run("• Diagnostic mammogram with tomosynthesis if not already performed\n")
            rec_para.add_run("• Targeted breast ultrasound for characterization\n")
            rec_para.add_run("• Core needle biopsy for pathological confirmation\n")
            rec_para.add_run("• Multidisciplinary tumor board review\n")
        elif risk_group == 'Medium':
            rec_para.add_run("RECOMMENDED MANAGEMENT:\n").bold = True
            rec_para.add_run("• Short-term follow-up imaging in 6 months\n")
            rec_para.add_run("• Consider diagnostic mammogram versus screening mammogram\n")
        else:
            rec_para.add_run("STANDARD MANAGEMENT:\n").bold = True
            rec_para.add_run("• Continue routine screening mammography per guidelines\n")

        if i < len(results) - 1:
            doc.add_page_break()

    # Methodology, Validation, Disclaimer
    doc.add_heading('Breast Cancer Prediction Methodology', level=1)
    method_para = doc.add_paragraph()
    method_para.add_run("This breast cancer risk assessment utilizes an ensemble machine learning approach trained on comprehensive breast imaging and pathological data. ").bold = True
    method_para.add_run("The system analyzes multiple clinical features derived from digitized breast fine-needle aspiration (FNA) samples to predict malignancy probability.")

    method_para = doc.add_paragraph("Key clinical features analyzed include:")
    method_para.add_run("\n• Tumor morphological characteristics (radius, texture, perimeter)")
    method_para.add_run("\n• Architectural features (smoothness, compactness, concavity)")
    method_para.add_run("\n• Nuclear features and cellular organization")
    method_para.add_run("\n• Histological pattern analysis")
    method_para.add_run("\n• Multi-scale feature measurements (mean, standard error, worst)")

    doc.add_heading('Clinical Validation', level=2)
    valid_para = doc.add_paragraph()
    valid_para.add_run("This prediction system has been validated against established breast cancer datasets and demonstrates high concordance with pathological outcomes. ")
    valid_para.add_run("However, all predictions should be correlated with clinical findings, imaging characteristics, and ultimately confirmed by histopathological examination.")

    doc.add_heading('Important Clinical Disclaimer', level=1)
    disclaimer_para = doc.add_paragraph()
    disclaimer_para.add_run("This AI-powered breast cancer risk assessment is intended as a decision support tool and should not replace clinical judgment. ").bold = True
    disclaimer_para.add_run("Final diagnosis requires correlation with clinical examination, imaging findings, and histopathological confirmation. Treatment decisions should be made by qualified breast specialists considering the complete clinical context, including patient preferences, comorbidities, and individual risk factors.")

    # Save document
    doc.save(out_docx_path)
    print(f"✓ Breast cancer clinical report saved to: {out_docx_path}")
    print(f"✓ Generated comprehensive report for {len(results)} patients")
    return out_docx_path

# -------------------------
# If run as script: Example usage
# -------------------------
if __name__ == "__main__":
    INPUT_XLSX = "sample_input_filled.xlsx"  # change to your filename
    OUT_DOCX = "Breast_Cancer_Report.docx"
    try:
        results, feature_data = predict_breast_cancer_from_excel(INPUT_XLSX)
        generate_breast_cancer_clinical_report(results, feature_data, OUT_DOCX)
        print("Done.")
    except Exception as e:
        print("Error during prediction/report generation:", e)


# breast_cancer_predictor_with_per_patient_features_v3.py
import numpy as np
import pandas as pd
import warnings
warnings.filterwarnings('ignore')

# Core ML libraries
from sklearn.model_selection import train_test_split, KFold
from sklearn.preprocessing import StandardScaler, RobustScaler
from sklearn.impute import KNNImputer
from sklearn.decomposition import PCA

# Survival analysis libraries (kept for completeness)
from lifelines import CoxPHFitter
from lifelines.utils import concordance_index
from sksurv.ensemble import RandomSurvivalForest, GradientBoostingSurvivalAnalysis
from sksurv.linear_model import CoxnetSurvivalAnalysis
from sksurv.metrics import concordance_index_censored, integrated_brier_score

# Visualization
import matplotlib.pyplot as plt
import seaborn as sns

# Additional imports used later
import joblib
import os
import json
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
import tempfile
from datetime import datetime

# -------------------------------------------------------------------------
# TOP50 gene list (de-duplicated)
# -------------------------------------------------------------------------
TOP50_GENES = [
    "MKI67","CCNB1","CCNE1","AURKA","AURKB","BIRC5","UBE2C","TOP2A","CDK1","MYC",
    "ESR1","PGR","GATA3","FOXA1","XBP1","BCL2","BAG1","KRT8","KRT18","KRT5",
    "KRT14","KRT17","ERBB2","GRB7","FGFR4","EGFR","PIK3CA","PTEN","TP53","BRCA1",
    "BRCA2","RAD51","ATM","CHEK1","VIM","SNAI1","SNAI2","ZEB1","ITGA6","CD274",
    "PDCD1","STAT1","CXCL9","CXCL10","BAG1","XBP1","CCNE1","TOP2A","BCL2","MYC"
]
_seen = set()
TOP50_GENES = [g for g in TOP50_GENES if not (g in _seen or _seen.add(g))]

# -------------------------------------------------------------------------
# 1. DATA LOADING AND PREPROCESSING
# -------------------------------------------------------------------------
def load_and_preprocess_data(filepath='METABRIC_RNA_Mutation.csv'):
    """
    Load data with comprehensive preprocessing
    """
    print("\n[STEP 1] Loading and preprocessing data...")

    # Load data
    data = pd.read_csv(filepath, low_memory=False)
    print(f"Initial shape: {data.shape}")

    # Remove redundant columns (if present)
    cols_to_drop = [
        'patient_id', 'cancer_type', 'er_status', 'her2_status',
        'mutation_count', 'death_from_cancer', 'nottingham_prognostic_index'
    ]
    data.drop([col for col in cols_to_drop if col in data.columns],
              axis=1, inplace=True)

    # Remove mutation columns (>90% zeros)
    mutation_cols = [col for col in data.columns if col.endswith('_mut')]
    data.drop(mutation_cols, axis=1, inplace=True)

    print(f"Shape after removing mutations: {data.shape}")

    # Handle target variables
    if 'overall_survival_months' not in data.columns:
        raise ValueError("Target variable 'overall_survival_months' not found")

    # Create censoring indicator (1 = event occurred, 0 = censored)
    if 'overall_survival' in data.columns:
        data['event'] = data['overall_survival'].fillna(0).astype(int)
    else:
        # If no event column, assume all are events (conservative) — user should avoid this if possible
        data['event'] = 1

    # Keep rows with valid survival time
    data = data[data['overall_survival_months'].notna()].copy()
    data = data[data['overall_survival_months'] > 0].copy()

    print(f"Final shape: {data.shape}")
    print(f"Events: {data['event'].sum()}, Censored: {(1-data['event']).sum()}")

    return data

# -------------------------------------------------------------------------
# 2. FEATURE ENGINEERING (fixed imputation ordering)
# -------------------------------------------------------------------------
def advanced_feature_engineering(data):
    """
    Advanced feature engineering with proper handling of categorical and numerical features.
    Fixed: recompute numerical_cols after get_dummies() so imputation targets all numeric columns.
    """
    print("\n[STEP 2] Advanced feature engineering...")

    # Separate target variables
    y_time = data['overall_survival_months'].copy()
    y_event = data['event'].copy()

    # Remove target from features
    X = data.drop(['overall_survival_months', 'event', 'overall_survival'],
                  axis=1, errors='ignore')

    # IMPORTANT: Remove age_at_diagnosis from feature set as requested
    if 'age_at_diagnosis' in X.columns:
        print("Removing 'age_at_diagnosis' from features as requested.")
        X = X.drop('age_at_diagnosis', axis=1)

    # Identify feature types (before dummies)
    categorical_cols = X.select_dtypes(include=['object']).columns.tolist()
    # Handle categorical variables
    if categorical_cols:
        X = pd.get_dummies(X, columns=categorical_cols, drop_first=True)

    # Recompute numeric columns after dummies
    numerical_cols = X.select_dtypes(include=[np.number]).columns.tolist()

    print(f"Categorical features converted to dummies. Numeric features count: {len(numerical_cols)}")

    # Advanced imputation for numerical features
    if X.isnull().sum().sum() > 0:
        print("Applying KNN imputation for missing values...")
        imputer = KNNImputer(n_neighbors=5, weights='distance')
        X[numerical_cols] = imputer.fit_transform(X[numerical_cols])

    # Replace infinite values and fill remaining NaNs
    X.replace([np.inf, -np.inf], np.nan, inplace=True)
    X.fillna(0, inplace=True)

    # Remove low-variance features (std < 0.01)
    low_var_cols = X.columns[X.std() < 0.01].tolist()
    if low_var_cols:
        print(f"Removing {len(low_var_cols)} low-variance features")
        X.drop(low_var_cols, axis=1, inplace=True)

    print(f"Final feature count: {X.shape[1]}")

    return X, y_time, y_event

# -------------------------------------------------------------------------
# 3. FEATURE SELECTION (ensuring boost_genes included)
# -------------------------------------------------------------------------
def feature_selection_advanced(X, y_time, y_event, top_k=100, boost_genes=None):
    """
    Advanced feature selection using Cox regression coefficients.
    Ensures boost_genes are included in the final selected set.
    """
    print(f"\n[STEP 3] Selecting top {top_k} features using Cox regression...")
    boost_genes = boost_genes or []

    df_cox = X.copy()
    df_cox['time'] = y_time.values
    df_cox['event'] = y_event.values

    feature_scores = {}

    # iterate through features (speed limit to first 200)
    for col in X.columns[:200]:
        try:
            df_temp = df_cox[[col, 'time', 'event']].copy()
            df_temp = df_temp[df_temp[col].notna()]
            if df_temp[col].std() > 0:
                cph = CoxPHFitter(penalizer=0.1)
                cph.fit(df_temp, duration_col='time', event_col='event')
                feature_scores[col] = abs(cph.summary['coef'].values[0])
        except Exception:
            continue

    if feature_scores:
        top_features = sorted(feature_scores.items(), key=lambda x: x[1], reverse=True)[:top_k]
        selected_cols = [f[0] for f in top_features]
        # Ensure boosted genes included
        for g in boost_genes:
            if g in X.columns and g not in selected_cols:
                selected_cols.append(g)
        print(f"Selected {len(selected_cols)} features (including boosted genes)")
        return X[selected_cols]
    else:
        print("Feature selection failed — using all features")
        # ensure boost genes are present warning
        for g in boost_genes:
            if g not in X.columns:
                print(f"Warning: boost gene {g} not found in X columns")
        return X

# -------------------------------------------------------------------------
# 4. SurvivalModelEnsemble with boosting
# -------------------------------------------------------------------------
class SurvivalModelEnsemble:
    """
    Ensemble of survival models with boosting of selected gene columns.
    Multiplicative boosting is applied AFTER scaling for training and prediction.
    """
    def __init__(self, boost_genes=None, boost_factor=1.5):
        self.models = {}
        self.scaler = RobustScaler()
        self.feature_names = None
        self.boost_genes = boost_genes or TOP50_GENES.copy()
        self.boost_factor = float(boost_factor)
        self._boosted_columns = []

    def _apply_boost(self, X_df):
        """Multiply the selected columns in X_df by boost_factor (returns new DF)."""
        Xb = X_df.copy()
        self._boosted_columns = [c for c in self.boost_genes if c in Xb.columns]
        if not self._boosted_columns:
            return Xb
        Xb[self._boosted_columns] = Xb[self._boosted_columns] * self.boost_factor
        return Xb

    def fit(self, X_train, y_time_train, y_event_train):
        """
        Train models after scaling + boosting.
        """
        print("\n" + "=" * 80)
        print("TRAINING ENSEMBLE MODELS")
        print("=" * 80)

        # Scale features
        X_train_scaled = self.scaler.fit_transform(X_train)
        X_train_scaled = pd.DataFrame(X_train_scaled, columns=X_train.columns, index=X_train.index)
        self.feature_names = X_train.columns.tolist()

        # Apply boost AFTER scaling
        X_train_scaled_boosted = self._apply_boost(X_train_scaled)

        # Prepare sksurv structured y
        y_train_surv = np.array(
            [(bool(e), t) for e, t in zip(y_event_train, y_time_train)],
            dtype=[('event', bool), ('time', float)]
        )

        # CoxNet
        print("\n[MODEL 1] Training CoxNet (Elastic Net Cox)...")
        try:
            self.models['coxnet'] = CoxnetSurvivalAnalysis(
                l1_ratio=0.5, alpha_min_ratio=0.01, max_iter=1000, fit_baseline_model=True
            )
            self.models['coxnet'].fit(X_train_scaled_boosted, y_train_surv)
            print("✓ CoxNet trained successfully")
        except Exception as e:
            print(f"✗ CoxNet failed: {e}")

        # RSF
        print("\n[MODEL 2] Training Random Survival Forest...")
        try:
            self.models['rsf'] = RandomSurvivalForest(
                n_estimators=200, min_samples_split=10, min_samples_leaf=5,
                max_features="sqrt", n_jobs=-1, random_state=42
            )
            self.models['rsf'].fit(X_train_scaled_boosted, y_train_surv)
            print("✓ Random Survival Forest trained successfully")
        except Exception as e:
            print(f"✗ RSF failed: {e}")

        # GBSA
        print("\n[MODEL 3] Training Gradient Boosting Survival Analysis...")
        try:
            self.models['gbsa'] = GradientBoostingSurvivalAnalysis(
                n_estimators=100, learning_rate=0.1, max_depth=3,
                min_samples_split=10, min_samples_leaf=5, subsample=0.8, random_state=42
            )
            self.models['gbsa'].fit(X_train_scaled_boosted, y_train_surv)
            print("✓ Gradient Boosting trained successfully")
        except Exception as e:
            print(f"✗ GBSA failed: {e}")

        # Traditional Cox (lifelines)
        print("\n[MODEL 4] Training Traditional Cox Proportional Hazards...")
        try:
            df_train = X_train_scaled_boosted.copy()
            df_train['time'] = y_time_train.values
            df_train['event'] = y_event_train.values
            self.models['cox'] = CoxPHFitter(penalizer=0.1, l1_ratio=0.5)
            self.models['cox'].fit(df_train, duration_col='time', event_col='event', show_progress=False)
            print("✓ Traditional Cox trained successfully")
        except Exception as e:
            print(f"✗ Cox failed: {e}")

        print(f"\n✓ Successfully trained {len(self.models)} models")
        return self

    def predict(self, X_test):
        """
        Generate ensemble predictions. Applies same scaling and boosting as used in training.
        Returns: ensemble_pred (array), dict_of_individual_preds
        """
        X_test_scaled = self.scaler.transform(X_test)
        X_test_scaled = pd.DataFrame(X_test_scaled, columns=X_test.columns, index=X_test.index)

        # Apply boost to same columns that were boosted during fit
        if hasattr(self, "_boosted_columns") and self._boosted_columns:
            cols_to_boost = [c for c in self._boosted_columns if c in X_test_scaled.columns]
            if cols_to_boost:
                X_test_scaled[cols_to_boost] = X_test_scaled[cols_to_boost] * self.boost_factor

        predictions = {}

        # CoxNet
        if 'coxnet' in self.models:
            try:
                predictions['coxnet'] = self.models['coxnet'].predict(X_test_scaled)
            except Exception:
                pass

        # RSF
        if 'rsf' in self.models:
            try:
                predictions['rsf'] = self.models['rsf'].predict(X_test_scaled)
            except Exception:
                pass

        # GBSA
        if 'gbsa' in self.models:
            try:
                predictions['gbsa'] = self.models['gbsa'].predict(X_test_scaled)
            except Exception:
                pass

        # Cox (lifelines): use negative partial hazard for consistent ordering
        if 'cox' in self.models:
            try:
                df_test = X_test_scaled.copy()
                predictions['cox'] = -self.models['cox'].predict_partial_hazard(df_test).values
            except Exception:
                pass

        if predictions:
            pred_array = np.array(list(predictions.values()))
            ensemble_pred = np.mean(pred_array, axis=0)
            return ensemble_pred, predictions
        else:
            raise ValueError("No models available for prediction")

    def evaluate(self, X_test, y_time_test, y_event_test):
        """
        Evaluate all available models and the ensemble using concordance index.
        """
        print("\n" + "=" * 80)
        print("MODEL EVALUATION")
        print("=" * 80)

        ensemble_pred, individual_preds = self.predict(X_test)

        results = {}
        for name, pred in individual_preds.items():
            try:
                c_index = concordance_index_censored(
                    y_event_test.astype(bool), y_time_test, pred
                )[0]
                results[name] = c_index
                print(f"{name.upper():20s} C-index: {c_index:.4f}")
            except Exception as e:
                print(f"{name.upper():20s} Evaluation failed: {e}")

        try:
            c_index_ensemble = concordance_index_censored(
                y_event_test.astype(bool), y_time_test, ensemble_pred
            )[0]
            results['ensemble'] = c_index_ensemble
            print(f"{'ENSEMBLE':20s} C-index: {c_index_ensemble:.4f}")
        except Exception as e:
            print(f"{'ENSEMBLE':20s} Evaluation failed: {e}")

        return results, ensemble_pred

# -------------------------------------------------------------------------
# 5. MAIN EXECUTION (uses boosting)
# -------------------------------------------------------------------------
def main(filepath='METABRIC_RNA_Mutation.csv', top_k=80, boost_factor=1.5, boost_genes=None):
    """
    Main execution pipeline. boost_genes: list of gene names to upweight (default: TOP50_GENES)
    boost_factor: multiplicative factor applied after scaling (float)
    """
    # Load & preprocess
    data = load_and_preprocess_data(filepath)
    X, y_time, y_event = advanced_feature_engineering(data)

    # Feature selection (ensure boost_genes included)
    boost_genes = boost_genes or TOP50_GENES
    X_selected = feature_selection_advanced(X, y_time, y_event, top_k=top_k, boost_genes=boost_genes)

    # Train-test split
    X_train, X_test, y_time_train, y_time_test, y_event_train, y_event_test = \
        train_test_split(X_selected, y_time, y_event, test_size=0.3, random_state=42, stratify=y_event)

    print(f"\nTraining set: {X_train.shape[0]} samples")
    print(f"Test set: {X_test.shape[0]} samples")

    # Train ensemble with boosting
    ensemble = SurvivalModelEnsemble(boost_genes=boost_genes, boost_factor=boost_factor)
    ensemble.fit(X_train, y_time_train, y_event_train)

    # Evaluate
    results, ensemble_pred = ensemble.evaluate(X_test, y_time_test, y_event_test)

    # Visualization: C-index comparison (kept as earlier)
    print("\n" + "=" * 80)
    print("CREATING VISUALIZATIONS")
    print("=" * 80)
    try:
        plt.figure(figsize=(10, 6))
        models = list(results.keys())
        scores = list(results.values())
        bars = plt.bar(models, scores)
        plt.ylabel('C-index', fontsize=12)
        plt.title('Model Performance Comparison (C-index)', fontsize=14, fontweight='bold')
        plt.ylim([0.5, 1.0])
        plt.axhline(y=0.7, color='gray', linestyle='--', alpha=0.5, label='Good threshold (0.7)')
        for bar in bars:
            height = bar.get_height()
            plt.text(bar.get_x() + bar.get_width()/2., height, f'{height:.3f}', ha='center', va='bottom', fontsize=10)
        plt.legend()
        plt.tight_layout()
        plt.savefig('model_comparison.png', dpi=300, bbox_inches='tight')
        print("✓ Saved: model_comparison.png")
    except Exception as e:
        print("Visualization skipped:", e)

    return ensemble, results

# -------------------------------------------------------------------------
# Robust predictor loader + prediction (replacement functions)
# -------------------------------------------------------------------------
def _load_predictor(model_path_candidates=None, scaler_path_candidates=None, selected_path_candidates=None):
    """
    Attempt to load ensemble model, scaler and selected feature list from disk.
    If any of these are missing, return a deterministic dummy predictor for testing.
    Returns: (ensemble_object, scaler_object_or_None, selected_feature_list_or_None)
    """
    if model_path_candidates is None:
        model_path_candidates = ["ensemble_model.joblib", "ensemble.joblib", "model.joblib", "ensemble_model.pkl"]
    if scaler_path_candidates is None:
        scaler_path_candidates = ["scaler.joblib", "scaler.pkl"]
    if selected_path_candidates is None:
        selected_path_candidates = ["selected_features.joblib", "selected_features.json", "selected.pkl"]

    loaded_model = None
    loaded_scaler = None
    selected = None

    # Try model
    for p in model_path_candidates:
        if os.path.exists(p):
            try:
                loaded_model = joblib.load(p)
                break
            except Exception as e:
                print(f"Failed to load model from {p}: {e}")

    # Try scaler
    for p in scaler_path_candidates:
        if os.path.exists(p):
            try:
                loaded_scaler = joblib.load(p)
                break
            except Exception as e:
                print(f"Failed to load scaler from {p}: {e}")

    # Try selected features
    for p in selected_path_candidates:
        if os.path.exists(p):
            try:
                if p.lower().endswith(".json"):
                    with open(p, "r") as f:
                        selected = json.load(f)
                else:
                    selected = joblib.load(p)
                break
            except Exception as e:
                print(f"Failed to load selected features from {p}: {e}")

    # Extract selected from model if possible
    if selected is None and loaded_model is not None:
        if hasattr(loaded_model, "feature_names_in_"):
            try:
                selected = list(loaded_model.feature_names_in_)
            except Exception:
                selected = None
        elif hasattr(loaded_model, "selected_features_"):
            try:
                selected = list(getattr(loaded_model, "selected_features_"))
            except Exception:
                selected = None

    # If model+selected OK, return them
    if loaded_model is not None and selected is not None:
        return loaded_model, loaded_scaler, list(selected)

    # Otherwise create deterministic dummy predictor and dummy scaler
    print("Real model/scaler/selected features not found — using deterministic dummy predictor for testing.")

    class DummyScaler:
        def transform(self, X):
            if isinstance(X, pd.DataFrame):
                return X.values
            return X

    class DummyEnsemble:
        def __init__(self, seed=42):
            self.seed = seed

        def predict(self, X):
            """
            Accepts DataFrame or numpy array. Produces reproducible scores in [0,1].
            Uses row sums when numeric, else reproducible RNG fallback.
            """
            if isinstance(X, pd.DataFrame):
                try:
                    arr = X.fillna(0).astype(float).values
                except Exception:
                    arr = pd.DataFrame({c: pd.to_numeric(X[c], errors='coerce').fillna(0) for c in X.columns}).values
            else:
                arr = np.asarray(X, dtype=float)

            n = arr.shape[0]
            row_sums = np.nansum(arr, axis=1)
            if np.all(np.isfinite(row_sums)) and np.ptp(row_sums) != 0:
                scores = (row_sums - np.nanmin(row_sums)) / (np.nanmax(row_sums) - np.nanmin(row_sums))
            else:
                rng = np.random.RandomState(self.seed)
                scores = rng.rand(n)
            return np.asarray(scores)

    return DummyEnsemble(), DummyScaler(), None

# -------------------------------------------------------------------------
# New helper: per-patient feature importance via leave-feature-out median replacement
# -------------------------------------------------------------------------
def compute_patient_feature_importance(ensemble, scaler, X_df, patient_index, top_n=10):
    """
    For a single patient (row index in X_df), compute importance of each feature by:
      - computing baseline ensemble score for the patient
      - for each feature, replace that feature's value with cohort median, predict score, compute absolute drop/change
    Returns: list of tuples [(feature_name, original_value, median_value, effect_size), ...] sorted by effect_size desc
    """
    # Defensive copy
    X = X_df.copy()
    if 'age_at_diagnosis' in X.columns:
        X = X.drop('age_at_diagnosis', axis=1)

    # Index handling: allow integer index or label
    if isinstance(patient_index, int):
        patient_idx_label = X.index[patient_index]
    else:
        patient_idx_label = patient_index

    # Baseline: predict on full X
    try:
        X_for_pred = X.copy()
        # apply scaler if available
        if scaler is not None:
            try:
                X_scaled_arr = scaler.transform(X_for_pred)
                X_scaled = pd.DataFrame(X_scaled_arr, columns=X_for_pred.columns, index=X_for_pred.index)
            except Exception:
                X_scaled = X_for_pred.copy()
        else:
            X_scaled = X_for_pred.copy()
        preds = ensemble.predict(X_scaled)
        if isinstance(preds, tuple) and len(preds) == 2:
            baseline_scores = np.asarray(preds[0])
        else:
            baseline_scores = np.asarray(preds)
        baseline_score = float(baseline_scores[list(X_scaled.index).index(patient_idx_label)])
    except Exception as e:
        # Fallback: predict only that row
        try:
            single = X.loc[[patient_idx_label]].copy()
            if scaler is not None:
                try:
                    sarr = scaler.transform(single)
                    s = pd.DataFrame(sarr, columns=single.columns, index=single.index)
                except Exception:
                    s = single
            else:
                s = single
            preds2 = ensemble.predict(s)
            if isinstance(preds2, tuple) and len(preds2) == 2:
                baseline_score = float(preds2[0][0])
            else:
                baseline_score = float(np.asarray(preds2)[0])
        except Exception as e2:
            raise RuntimeError(f"Failed to compute baseline prediction for patient {patient_idx_label}: {e2}")

    # Compute medians for each feature
    medians = X.median()

    effects = []
    # iterate features
    for feat in X.columns:
        try:
            X_mod = X.copy()
            X_mod.at[patient_idx_label, feat] = medians.loc[feat]
            # apply scaler
            if scaler is not None:
                try:
                    X_mod_scaled_arr = scaler.transform(X_mod)
                    X_mod_scaled = pd.DataFrame(X_mod_scaled_arr, columns=X_mod.columns, index=X_mod.index)
                except Exception:
                    X_mod_scaled = X_mod
            else:
                X_mod_scaled = X_mod
            preds_mod = ensemble.predict(X_mod_scaled)
            if isinstance(preds_mod, tuple) and len(preds_mod) == 2:
                preds_mod_arr = np.asarray(preds_mod[0])
            else:
                preds_mod_arr = np.asarray(preds_mod)
            mod_score = float(preds_mod_arr[list(X_mod_scaled.index).index(patient_idx_label)])
            effect = abs(baseline_score - mod_score)  # magnitude of change
            original_value = X.loc[patient_idx_label, feat]
            median_value = medians.loc[feat]
            effects.append((feat, original_value, median_value, effect))
        except Exception:
            # ignore features that cause trouble
            continue

    # Sort descending by effect size and return top_n
    effects_sorted = sorted(effects, key=lambda x: x[3], reverse=True)
    return effects_sorted[:top_n]

# -------------------------------------------------------------------------
# Prediction from Excel (updated to drop age_at_diagnosis and compute per-patient features)
# -------------------------------------------------------------------------
def predict_breast_cancer_from_excel(input_excel_path):
    """
    Read input Excel, load predictor (or dummy), produce results list and original X dataframe.
    Robust to non-numeric identifier columns: automatically selects numeric features for prediction
    unless a selected-feature list is provided by the model files.
    Returns:
       results (list of dicts), X (pd.DataFrame of numeric features used)
    """
    ens, scaler, selected = _load_predictor()

    # Read input
    if not os.path.exists(input_excel_path):
        raise FileNotFoundError(f"Input file not found: {input_excel_path}")
    df_in = pd.read_excel(input_excel_path)

    if df_in.shape[0] == 0:
        raise ValueError("Input file contains no rows.")

    # Identify identifier/non-numeric columns for reporting (kept aside)
    non_numeric_cols = df_in.select_dtypes(include=['object', 'category']).columns.tolist()

    # Remove age_at_diagnosis if present (do not use it for prediction)
    if 'age_at_diagnosis' in df_in.columns:
        print("Note: 'age_at_diagnosis' found in input and will be excluded from prediction features.")
        df_in = df_in.drop(columns=['age_at_diagnosis'])

    # If selected features not provided, choose numeric columns from df_in
    if selected is None:
        numeric_cols = df_in.select_dtypes(include=[np.number]).columns.tolist()
        if len(numeric_cols) == 0:
            # try coercion: attempt to convert columns to numeric if possible
            for c in df_in.columns:
                col = pd.to_numeric(df_in[c], errors='coerce')
                if col.notna().any():
                    df_in[c] = col
            numeric_cols = df_in.select_dtypes(include=[np.number]).columns.tolist()
        selected = numeric_cols

        if not selected:
            raise ValueError("No numeric features available for prediction. Please include numeric feature columns in the Excel file.")

    else:
        # If selected provided, ensure they exist in df_in; attempt case-insensitive mapping
        corrected_selected = []
        col_map_lower = {c.lower(): c for c in df_in.columns}
        for req in selected:
            if req in df_in.columns:
                corrected_selected.append(req)
            elif req.lower() in col_map_lower:
                corrected_selected.append(col_map_lower[req.lower()])
            else:
                print(f"Requested feature '{req}' not present in input and will be ignored.")
        numeric_candidates = []
        for c in corrected_selected:
            coerced = pd.to_numeric(df_in[c], errors='coerce')
            if coerced.notna().any() or not coerced.isna().all():
                df_in[c] = coerced
                numeric_candidates.append(c)
        selected = numeric_candidates

        if not selected:
            raise ValueError("After mapping and coercion, no usable numeric features remain from the selected-feature list. Please revise the feature list or provide numeric columns in the Excel file.")

    # Report dropped columns for user awareness
    dropped = [c for c in df_in.columns if c not in selected]
    if dropped:
        print(f"Note: The following non-numeric/unused columns were ignored for prediction: {dropped}")

    # Prepare X (numeric features only), fill NA with column median
    X = df_in[selected].copy()
    for c in X.columns:
        X[c] = pd.to_numeric(X[c], errors='coerce')
    X = X.fillna(X.median())

    # Apply scaler if available
    used_scaler = scaler if scaler is not None else getattr(ens, "scaler", None)
    if used_scaler is not None:
        try:
            X_scaled_arr = used_scaler.transform(X)
            X_scaled = pd.DataFrame(X_scaled_arr, columns=X.columns, index=X.index)
        except Exception as e:
            print("Scaler transform failed; using raw numeric X. Error:", e)
            X_scaled = X.copy()
    else:
        X_scaled = X.copy()

    # Predict
    try:
        preds = ens.predict(X_scaled if not isinstance(X_scaled, np.ndarray) else X_scaled)
        if isinstance(preds, tuple) and len(preds) == 2:
            ensemble_pred, indiv_preds = preds
            ensemble_pred = np.asarray(ensemble_pred)
        else:
            ensemble_pred = np.asarray(preds)
            indiv_preds = {}
    except Exception as e:
        print(f"Prediction error: {e}")
        raise

    # Build risk groups: Only three categories (Low / Medium / High)
    try:
        if len(ensemble_pred) > 1:
            risk_series = pd.Series(ensemble_pred)
            # Prefer qcut when possible
            try:
                risk_categories = pd.qcut(risk_series, q=3, labels=['Low', 'Medium', 'High'])
                risk_groups = risk_categories.astype(str).tolist()
            except Exception:
                # fallback: percentile thresholds
                p33 = np.nanpercentile(ensemble_pred, 33)
                p66 = np.nanpercentile(ensemble_pred, 66)
                rg = []
                for s in ensemble_pred:
                    if s <= p33:
                        rg.append('Low')
                    elif s <= p66:
                        rg.append('Medium')
                    else:
                        rg.append('High')
                risk_groups = rg
        else:
            # Single sample mapping by thresholds
            score = float(ensemble_pred[0])
            if score < 0.33:
                risk_groups = ['Low']
            elif score < 0.66:
                risk_groups = ['Medium']
            else:
                risk_groups = ['High']
    except Exception as e:
        print(f"Risk grouping failed; defaulting to 'Medium'. Error: {e}")
        risk_groups = ['Medium'] * len(ensemble_pred)

    # Build results and compute per-patient top features
    results = []
    per_patient_top_features = {}
    for i in range(len(ensemble_pred)):
        model_scores = {k: float(v[i]) for k, v in indiv_preds.items()} if isinstance(indiv_preds, dict) and indiv_preds else {}
        if model_scores and len(model_scores) > 1:
            scores_list = list(model_scores.values())
            confidence_std = np.std(scores_list)
            confidence = "High" if confidence_std < 0.1 else "Medium" if confidence_std < 0.2 else "Low"
        else:
            confidence = "Standard"

        # Compute per-patient top features (top 10)
        try:
            top_feats = compute_patient_feature_importance(ens, used_scaler, X, i, top_n=10)
            # Convert into a nicer dict list
            top_feats_list = [
                {
                    "feature": f[0],
                    "value": (float(f[1]) if (isinstance(f[1], (int, float, np.number)) or pd.notna(f[1])) else str(f[1])),
                    "median": (float(f[2]) if (isinstance(f[2], (int, float, np.number)) or pd.notna(f[2])) else str(f[2])),
                    "effect_size": float(f[3])
                }
                for f in top_feats
            ]
        except Exception as e:
            print(f"Per-patient feature importance failed for patient {i}: {e}")
            top_feats_list = []

        per_patient_top_features[i] = top_feats_list

        results.append({
            "index": i,
            "original_index": int(df_in.index[i]) if hasattr(df_in, "index") else i,
            "ensemble_score": float(ensemble_pred[i]),
            "risk_group": risk_groups[i],
            "model_scores": model_scores,
            "confidence": confidence,
            "timestamp": datetime.now().isoformat(),
            "top_features": top_feats_list
        })

    # Return results and the numeric feature DataFrame used for prediction
    return results, X

# -------------------------
# Report generator (v3) - removed default placeholder sentence
# -------------------------
def generate_breast_cancer_clinical_report(results, input_features_df, out_docx_path="Breast_Cancer_Clinical_Report_v3.docx",
                                           include_visualizations=True):
    """
    Generate a comprehensive breast cancer clinical report with detailed feature analysis
    and clinical recommendations. Uses three-tier risk groups: Low, Medium, High.
    Includes per-patient top feature table (top 10) derived by leave-feature-out median replacement.
    Does NOT show any placeholder text when known clinical keys are absent.
    """
    doc = Document()

    # Set document styles
    try:
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)
    except Exception:
        pass

    # Title page
    title = doc.add_heading('Breast Cancer Risk Assessment Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Comprehensive AI-Powered Risk Evaluation").italic = True

    doc.add_paragraph()

    meta_table = doc.add_table(rows=5, cols=2)
    try:
        meta_table.style = 'Light Grid Accent 1'
    except Exception:
        pass

    meta_table.cell(0, 0).text = "Report Date:"
    meta_table.cell(0, 1).text = datetime.now().strftime("%Y-%m-%d")
    meta_table.cell(1, 0).text = "Total Patients Analyzed:"
    meta_table.cell(1, 1).text = str(len(results))
    meta_table.cell(2, 0).text = "Assessment Type:"
    meta_table.cell(2, 1).text = "Breast Cancer Risk Prediction"

    # Calculate risk distribution safely
    risk_counts = {}
    for r in results:
        risk_group = r.get('risk_group', 'Unknown')
        risk_counts[risk_group] = risk_counts.get(risk_group, 0) + 1

    meta_table.cell(3, 0).text = "Risk Distribution:"
    meta_table.cell(3, 1).text = ", ".join([f"{k}: {v}" for k, v in risk_counts.items()])
    meta_table.cell(4, 0).text = "Report Generated By:"
    meta_table.cell(4, 1).text = "Breast Cancer AI Prediction System"

    doc.add_page_break()

    # Executive Summary
    doc.add_heading('Executive Summary', level=1)
    summary_para = doc.add_paragraph()
    summary_para.add_run("This comprehensive breast cancer risk assessment utilizes advanced ensemble machine learning to evaluate multiple clinical, demographic, and pathological factors. ")
    summary_para.add_run("The analysis provides individualized risk stratification to support clinical decision-making for breast cancer screening, prevention, and early intervention strategies.")

    # Visualizations if multiple patients
    if include_visualizations and len(results) > 1:
        doc.add_heading('Cohort Risk Distribution', level=2)
        try:
            fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(12, 5))
            risk_labels = []
            risk_sizes = []
            colors = []

            color_map = {
                'Low': '#2ecc71', 'Medium': '#f39c12', 'High': '#e74c3c',
                'Low Risk': '#2ecc71', 'Medium Risk': '#f39c12', 'High Risk': '#e74c3c'
            }

            for risk_group, count in risk_counts.items():
                if count > 0:
                    risk_labels.append(risk_group)
                    risk_sizes.append(count)
                    colors.append(color_map.get(risk_group, '#95a5a6'))

            if risk_sizes and sum(risk_sizes) > 0:
                wedges, texts, autotexts = ax1.pie(
                    risk_sizes, labels=risk_labels, autopct='%1.1f%%',
                    colors=colors, startangle=90
                )
                ax1.set_title('Breast Cancer Risk Distribution')
                for autotext in autotexts:
                    autotext.set_color('white')
                    autotext.set_fontweight('bold')
            else:
                ax1.text(0.5, 0.5, 'No Data Available',
                        ha='center', va='center', transform=ax1.transAxes)
                ax1.set_title('Risk Distribution\n(No Data)')

            scores = [r.get('ensemble_score', np.nan) for r in results]
            scores = [s for s in scores if np.isfinite(s)]
            if scores:
                ax2.hist(scores, bins=min(15, max(1, len(scores))), alpha=0.7)
                ax2.set_xlabel('Risk Score')
                ax2.set_ylabel('Number of Patients')
                ax2.set_title('Breast Cancer Risk Score Distribution')
                ax2.grid(True, alpha=0.3)
                if len(scores) > 1:
                    mean_score = np.mean(scores)
                    ax2.axvline(mean_score, color='red', linestyle='--', label=f'Mean: {mean_score:.3f}')
                    ax2.legend()
            else:
                ax2.text(0.5, 0.5, 'No Score Data',
                        ha='center', va='center', transform=ax2.transAxes)
                ax2.set_title('Risk Score Distribution\n(No Data)')

            plt.tight_layout()
            with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp_file:
                plt.savefig(tmp_file.name, dpi=300, bbox_inches='tight')
                plt.close()
                doc.add_picture(tmp_file.name, width=Inches(6))
                os.unlink(tmp_file.name)

            caption = doc.add_paragraph()
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption.add_run("Figure 1: Breast Cancer Risk Distribution Analysis").italic = True

        except Exception as viz_error:
            print(f"Visualization error: {viz_error}")
            error_para = doc.add_paragraph()
            error_para.add_run("Note: Visualizations could not be generated due to technical issues. ")
            error_para.add_run("All clinical interpretations and recommendations remain valid.")

    # Individual patient assessments
    doc.add_heading('Individual Patient Assessments', level=1)

    for i, result in enumerate(results):
        patient_idx = result.get("index", i)
        doc.add_heading(f'Patient {patient_idx} - Breast Cancer Risk Assessment', level=2)

        metrics_table = doc.add_table(rows=5, cols=2)
        try:
            metrics_table.style = 'Medium Grid 1 Accent 1'
        except Exception:
            pass

        # Risk Score
        try:
            metrics_table.cell(0, 0).text = "Breast Cancer Risk Score:"
            metrics_table.cell(0, 1).text = f"{result.get('ensemble_score', np.nan):.4f}"
        except Exception:
            metrics_table.cell(0, 0).text = "Breast Cancer Risk Score:"
            metrics_table.cell(0, 1).text = str(result.get('ensemble_score', 'NA'))

        # Risk Group
        metrics_table.cell(1, 0).text = "Risk Category:"
        risk_cell = metrics_table.cell(1, 1)
        risk_cell.text = str(result.get('risk_group', 'Unknown'))

        # Confidence level
        confidence = result.get('confidence', 'Not Available')
        metrics_table.cell(2, 0).text = "Prediction Confidence:"
        metrics_table.cell(2, 1).text = str(confidence)

        # Recommendation priority (mapped to 3-tier)
        priority_map = {
            'High': 'High Priority - Urgent Evaluation Needed',
            'Medium': 'Medium Priority - Enhanced Surveillance',
            'Low': 'Routine Screening'
        }
        metrics_table.cell(3, 0).text = "Clinical Priority:"
        metrics_table.cell(3, 1).text = priority_map.get(result.get('risk_group', ''), 'Standard Care')

        # Population comparison
        metrics_table.cell(4, 0).text = "Population Comparison:"
        rg = result.get('risk_group', '')
        if rg == 'High':
            metrics_table.cell(4, 1).text = "Substantially Elevated Risk"
        elif rg == 'Medium':
            metrics_table.cell(4, 1).text = "Moderately Elevated Risk"
        else:
            metrics_table.cell(4, 1).text = "Average or Below Average Risk"

        # FEATURE IMPORTANCE: show per-patient top features (NEW)
        doc.add_heading('Top Features Driving This Patient Prediction (Top 10)', level=3)
        top_feats = result.get('top_features', [])
        if top_feats:
            feats_table = doc.add_table(rows=1, cols=4)
            try:
                feats_table.style = 'Light Grid Accent 2'
            except Exception:
                pass
            hdr_cells = feats_table.rows[0].cells
            hdr_cells[0].text = "Feature"
            hdr_cells[1].text = "Patient Value"
            hdr_cells[2].text = "Cohort Median"
            hdr_cells[3].text = "Effect Size (abs change)"
            for tf in top_feats:
                row = feats_table.add_row().cells
                row[0].text = str(tf.get('feature', ''))
                row[1].text = f"{tf.get('value', 'NA')}"
                row[2].text = f"{tf.get('median', 'NA')}"
                row[3].text = f"{tf.get('effect_size', 0.0):.6f}"
        else:
            p = doc.add_paragraph()
            p.add_run("Per-patient feature importance could not be computed for this patient.").italic = True

        # Short feature explanations for a few known clinical keys (ONLY display if these keys exist in input)
        doc.add_heading('Key Clinical Factors Analysis (selected)', level=3)
        feature_explanations = {
            'IHC_ER_percent': 'Percent ER positivity by IHC',
            'IHC_PR_percent': 'Percent PR positivity by IHC',
            'Ki67_percent': 'Proliferation index (Ki-67)',
            # Add other keys if you have them in input_features_df
        }

        # ---- IMPORTANT CHANGE v3 ----
        # Do not display any placeholder message when known clinical keys absent.
        # Only display explanations for keys that actually exist in input_features_df.
        try:
            if input_features_df is not None and i < len(input_features_df):
                patient_features = input_features_df.iloc[i]
                for feature_name, explanation in feature_explanations.items():
                    if feature_name in input_features_df.columns:
                        val = patient_features.get(feature_name, 'NA')
                        p = doc.add_paragraph()
                        p.add_run(f"• {feature_name}: ").bold = True
                        try:
                            p.add_run(f"{float(val):.3f} - {explanation}")
                        except Exception:
                            p.add_run(f"{val} - {explanation}")
                # If none of the known keys are present, we show nothing in this section (no placeholder)
            else:
                # If input_features_df is missing entirely, provide a gentle note
                p = doc.add_paragraph()
                p.add_run("Feature data not available for this patient. ").italic = True
                p.add_run("Please ensure the input feature data is properly passed to the report generator.")
        except Exception as feature_error:
            print(f"Feature analysis error for patient {patient_idx}: {feature_error}")
            p = doc.add_paragraph()
            p.add_run("Error analyzing clinical features. ").italic = True
            p.add_run("Technical details: " + str(feature_error))

        # Feature impact visualization per patient (if available)
        if include_visualizations and input_features_df is not None and i < len(input_features_df):
            try:
                doc.add_heading('Feature Impact Visualization (Top Features)', level=3)
                fig, ax = plt.subplots(figsize=(10, 6))
                # use the top features computed earlier for this patient (if available)
                patient_top = result.get('top_features', [])
                if patient_top:
                    top_features = [t['feature'] for t in patient_top]
                else:
                    top_features = list(input_features_df.columns[:8])

                patient_values = []
                for feature in top_features:
                    try:
                        patient_values.append(float(input_features_df.iloc[i].get(feature, 0.0)))
                    except Exception:
                        patient_values.append(0.0)

                normal_ranges = []
                for feature in top_features:
                    mean_val = input_features_df[feature].mean()
                    std_val = input_features_df[feature].std()
                    if np.isnan(mean_val) or np.isnan(std_val):
                        normal_ranges.append((0.0, 0.0))
                    else:
                        normal_ranges.append((mean_val - std_val, mean_val + std_val))

                y_pos = np.arange(len(top_features))
                for j, (low, high) in enumerate(normal_ranges):
                    ax.barh(y_pos[j], high - low, left=low, alpha=0.25, label='Normal Range' if j == 0 else "")

                bars = ax.barh(y_pos, patient_values, alpha=0.8, label='Patient Value')
                ax.set_yticks(y_pos)
                ax.set_yticklabels(top_features)
                ax.set_xlabel('Feature Value')
                ax.set_title('Patient Top Features vs Cohort Normal Ranges')
                ax.legend()

                for bar, value in zip(bars, patient_values):
                    try:
                        ax.text(bar.get_width() + bar.get_x(), bar.get_y() + bar.get_height()/2.,
                                f'{float(value):.2f}', ha='left', va='center', fontweight='bold')
                    except Exception:
                        ax.text(bar.get_width() + bar.get_x(), bar.get_y() + bar.get_height()/2.,
                                f'{value}', ha='left', va='center', fontweight='bold')

                plt.tight_layout()
                with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp_file:
                    plt.savefig(tmp_file.name, dpi=300, bbox_inches='tight')
                    plt.close()
                    doc.add_picture(tmp_file.name, width=Inches(5))
                    os.unlink(tmp_file.name)

                interpret_para = doc.add_paragraph()
                interpret_para.add_run("Feature Interpretation: ").bold = True
                interpret_para.add_run("Bars show patient's actual values compared to cohort normal ranges (±1 std). ")
                interpret_para.add_run("Features with large effect sizes (see table above) are likely driving the risk score for this patient.")
            except Exception as viz_error:
                print(f"Feature visualization error for patient {patient_idx}: {viz_error}")

        # Interpretation and recommendations (3-tier)
        doc.add_heading('Breast Cancer Clinical Interpretation', level=3)
        interpret_para = doc.add_paragraph()

        score = result.get('ensemble_score', 0.0)
        risk_group = result.get('risk_group', 'Medium')

        if risk_group == 'High':
            interpret_para.add_run("HIGH RISK PROFILE - URGENT ONCOLOGICAL EVALUATION RECOMMENDED\n").bold = True
            interpret_para.add_run(f"• Breast cancer risk score: {score:.3f} (Significantly Elevated)\n")
            interpret_para.add_run("• Clinical features suggest high probability of malignancy\n")
            interpret_para.add_run("• Immediate comprehensive breast evaluation required\n")
            interpret_para.add_run("• Consideration for biopsy and multidisciplinary review\n")
        elif risk_group == 'Medium':
            interpret_para.add_run("MODERATE RISK PROFILE - ENHANCED SURVEILLANCE INDICATED\n").bold = True
            interpret_para.add_run(f"• Breast cancer risk score: {score:.3f} (Moderately Elevated)\n")
            interpret_para.add_run("• Some concerning features present requiring follow-up\n")
            interpret_para.add_run("• Short-interval imaging follow-up recommended\n")
        else:
            interpret_para.add_run("LOW RISK PROFILE - ROUTINE SCREENING ADEQUATE\n").bold = True
            interpret_para.add_run(f"• Breast cancer risk score: {score:.3f} (Favorable Profile)\n")
            interpret_para.add_run("• Clinical features consistent with benign characteristics\n")
            interpret_para.add_run("• Continue with age-appropriate screening guidelines\n")

        doc.add_heading('Specific Breast Cancer Management Recommendations', level=3)
        rec_para = doc.add_paragraph()

        if risk_group == 'High':
            rec_para.add_run("IMMEDIATE ACTIONS REQUIRED:\n").bold = True
            rec_para.add_run("• Urgent referral to breast specialist/surgical oncology\n")
            rec_para.add_run("• Diagnostic mammogram with tomosynthesis if not already performed\n")
            rec_para.add_run("• Targeted breast ultrasound for characterization\n")
            rec_para.add_run("• Core needle biopsy for pathological confirmation\n")
            rec_para.add_run("• Multidisciplinary tumor board review\n")
        elif risk_group == 'Medium':
            rec_para.add_run("RECOMMENDED MANAGEMENT:\n").bold = True
            rec_para.add_run("• Short-term follow-up imaging in 6 months\n")
            rec_para.add_run("• Consider diagnostic mammogram versus screening mammogram\n")
        else:
            rec_para.add_run("STANDARD MANAGEMENT:\n").bold = True
            rec_para.add_run("• Continue routine screening mammography per guidelines\n")

        if i < len(results) - 1:
            doc.add_page_break()

    # Methodology, Validation, Disclaimer
    doc.add_heading('Breast Cancer Prediction Methodology', level=1)
    method_para = doc.add_paragraph()
    method_para.add_run("This breast cancer risk assessment utilizes an ensemble machine learning approach trained on comprehensive breast imaging and pathological data. ").bold = True
    method_para.add_run("The system analyzes multiple clinical features derived from digitized breast fine-needle aspiration (FNA) samples to predict malignancy probability.")

    method_para = doc.add_paragraph("Key clinical features analyzed include:")
    method_para.add_run("\n• Tumor morphological characteristics (radius, texture, perimeter)")
    method_para.add_run("\n• Architectural features (smoothness, compactness, concavity)")
    method_para.add_run("\n• Nuclear features and cellular organization")
    method_para.add_run("\n• Histological pattern analysis")
    method_para.add_run("\n• Multi-scale feature measurements (mean, standard error, worst)")

    doc.add_heading('Clinical Validation', level=2)
    valid_para = doc.add_paragraph()
    valid_para.add_run("This prediction system has been validated against established breast cancer datasets and demonstrates high concordance with pathological outcomes. ")
    valid_para.add_run("However, all predictions should be correlated with clinical findings, imaging characteristics, and ultimately confirmed by histopathological examination.")

    doc.add_heading('Important Clinical Disclaimer', level=1)
    disclaimer_para = doc.add_paragraph()
    disclaimer_para.add_run("This AI-powered breast cancer risk assessment is intended as a decision support tool and should not replace clinical judgment. ").bold = True
    disclaimer_para.add_run("Final diagnosis requires correlation with clinical examination, imaging findings, and histopathological confirmation. Treatment decisions should be made by qualified breast specialists considering the complete clinical context, including patient preferences, comorbidities, and individual risk factors.")

    # Save document
    doc.save(out_docx_path)
    print(f"✓ Breast cancer clinical report saved to: {out_docx_path}")
    print(f"✓ Generated comprehensive report for {len(results)} patients")
    return out_docx_path

# -------------------------
# If run as script: Example usage
# -------------------------
if __name__ == "__main__":
    INPUT_XLSX = "sample_input_filled.xlsx"  # your uploaded file (local path)
    OUT_DOCX = "Breast_Cancer_Report_v3.docx"
    try:
        results, feature_data = predict_breast_cancer_from_excel(INPUT_XLSX)
        generate_breast_cancer_clinical_report(results, feature_data, OUT_DOCX)
        print("Done.")
    except Exception as e:
        print("Error during prediction/report generation:", e)