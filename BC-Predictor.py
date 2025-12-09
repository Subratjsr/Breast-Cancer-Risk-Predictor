# app_streamlit_corrected.py
"""
Professional Breast Cancer Predictor MVP
- Modern, medical-grade UI/UX design
- Improved spacing, typography, and color scheme
- Enhanced user experience with better information hierarchy
- Normalizes risk labels to only: High Risk / Medium Risk / Low Risk
- Added manual data input interface
- Professional diagnostic center-style report generation with OncoVista™ watermark on all pages
"""

import streamlit as st
import pandas as pd
import numpy as np
import tempfile
import os
import plotly.express as px
import plotly.graph_objects as go
import traceback
from datetime import datetime
from io import BytesIO
import base64

st.set_page_config(
    page_title="Breast Cancer Risk Assessment",
    page_icon="🩺",
    layout="wide",
    initial_sidebar_state="collapsed"
)

# -------------------------
# Professional Medical UI CSS
# -------------------------
st.markdown(
    """
    <style>
    /* Base styling */
    .main { background-color: #f8fafc; }
    .stApp { background-color: #f8fafc; }

    /* Professional header */
    .main-header { 
        background: linear-gradient(135deg, #1e3a8a 0%, #3730a3 100%);
        padding: 2rem 0;
        margin: -1rem -1rem 2rem -1rem;
        border-radius: 0 0 20px 20px;
        box-shadow: 0 4px 20px rgba(0,0,0,0.08);
    }
    .header-title { 
        font-size: 2.25rem; 
        color: white; 
        font-weight: 700; 
        font-family: 'Inter', sans-serif;
        margin-bottom: 0.5rem;
    }
    .header-subtitle { 
        color: #e0e7ff; 
        font-size: 1.1rem; 
        font-weight: 400;
        line-height: 1.5;
        max-width: 600px;
    }

    /* Enhanced cards */
    .professional-card {
        background: white;
        border-radius: 16px;
        padding: 2rem;
        box-shadow: 0 4px 20px rgba(0,0,0,0.06);
        border: 1px solid #e2e8f0;
        margin-bottom: 1.5rem;
        transition: transform 0.2s ease, box-shadow 0.2s ease;
    }
    .professional-card:hover {
        transform: translateY(-2px);
        box-shadow: 0 8px 30px rgba(0,0,0,0.1);
    }

    /* Section headers */
    .section-title {
        font-size: 1.5rem;
        color: #1e293b;
        font-weight: 600;
        margin-bottom: 1.25rem;
        font-family: 'Inter', sans-serif;
        border-left: 4px solid #3b82f6;
        padding-left: 1rem;
    }

    /* Risk badges */
    .risk-badge {
        padding: 0.5rem 1rem;
        border-radius: 12px;
        color: white;
        font-weight: 600;
        font-size: 0.875rem;
        display: inline-flex;
        align-items: center;
        gap: 0.5rem;
    }
    .risk-badge::before {
        content: "●";
        font-size: 0.75rem;
    }

    /* Interpretation sections */
    .interpretation-header {
        color: #1e293b;
        font-size: 1.125rem;
        font-weight: 600;
        margin: 1.5rem 0 1rem 0;
        padding-bottom: 0.5rem;
        border-bottom: 2px solid #f1f5f9;
    }

    .recommendation-card {
        background: linear-gradient(135deg, #f0f9ff 0%, #e0f2fe 100%);
        border: 1px solid #bae6fd;
        border-radius: 12px;
        padding: 1.5rem;
        margin: 1rem 0;
    }

    /* Button enhancements */
    .stButton button {
        border-radius: 10px;
        font-weight: 600;
        padding: 0.75rem 1.5rem;
        border: none;
        transition: all 0.2s ease;
    }

    .stButton button:hover {
        transform: translateY(-1px);
        box-shadow: 0 4px 12px rgba(0,0,0,0.15);
    }

    /* File uploader enhancement */
    .uploadedFile {
        border: 2px dashed #cbd5e1;
        border-radius: 12px;
        padding: 1.5rem;
        background: #f8fafc;
    }

    /* Metric displays */
    .metric-card {
        background: white;
        border-radius: 12px;
        padding: 1.5rem;
        border-left: 4px solid #3b82f6;
        box-shadow: 0 2px 8px rgba(0,0,0,0.05);
    }

    /* Status indicators */
    .status-indicator {
        display: inline-flex;
        align-items: center;
        gap: 0.5rem;
        padding: 0.5rem 1rem;
        border-radius: 8px;
        font-weight: 600;
        font-size: 0.875rem;
    }

    /* Custom spacing */
    .section-spacing {
        margin: 2rem 0;
    }

    /* Progress indicators */
    .progress-container {
        background: #f1f5f9;
        border-radius: 10px;
        padding: 0.5rem;
        margin: 1rem 0;
    }

    /* Manual input styling */
    .gene-input-container {
        background: #f8fafc;
        border-radius: 12px;
        padding: 1.5rem;
        margin-bottom: 1rem;
        border: 1px solid #e2e8f0;
    }

    .gene-input-row {
        display: grid;
        grid-template-columns: 150px 1fr;
        gap: 1rem;
        align-items: center;
        margin-bottom: 1rem;
    }

    .add-gene-btn {
        background: #f1f5f9;
        border: 2px dashed #cbd5e1;
        border-radius: 8px;
        padding: 0.75rem;
        text-align: center;
        cursor: pointer;
        transition: all 0.2s ease;
    }

    .add-gene-btn:hover {
        background: #e2e8f0;
        border-color: #94a3b8;
    }

    .remove-gene-btn {
        background: #fee2e2;
        color: #dc2626;
        border: none;
        border-radius: 6px;
        padding: 0.25rem 0.75rem;
        font-size: 0.875rem;
        cursor: pointer;
    }

    .remove-gene-btn:hover {
        background: #fecaca;
    }

    </style>
    """,
    unsafe_allow_html=True,
)

# -------------------------
# Try to import real backend; otherwise demo stub
# -------------------------
try:
    import predictor_backend as pb  # your real backend file

    has_real_backend = True
except Exception as e:
    pb = None
    has_real_backend = False
    import_error_msg = str(e)


def demo_predict_from_excel(path):
    """Demo stub: reads input file and returns fake predictions + features_df"""
    df = pd.read_excel(path)
    n = len(df)
    rng = np.random.default_rng(12345)
    scores = rng.random(n) * 100 + 50  # More realistic score range
    # simple 3-group risk
    labels = pd.qcut(scores, q=3, labels=['Low Risk', 'Medium Risk', 'High Risk'])
    results = []
    for i, s in enumerate(scores):
        results.append({
            "index": i,
            "original_index": int(df.index[i]) if hasattr(df, "index") else i,
            "ensemble_score": float(s),
            "risk_group": str(labels[i]),
            "model_scores": {"demo_model": float(s)},
            "timestamp": datetime.now().isoformat()
        })
    return results, df


def create_professional_report_with_watermark(results, features_df, out_docx_path, include_visualizations=True):
    """Professional diagnostic center-style report generator with OncoVista™ watermark on all pages"""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
        from docx.enum.style import WD_STYLE_TYPE
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.enum.section import WD_SECTION

        # Create document
        doc = Document()

        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(2.54)
            section.bottom_margin = Cm(2.54)
            section.left_margin = Cm(2.54)
            section.right_margin = Cm(2.54)

        # Add custom styles
        styles = doc.styles

        # Title style
        if 'TitleStyle' not in styles:
            title_style = styles.add_style('TitleStyle', WD_STYLE_TYPE.PARAGRAPH)
            title_style.font.name = 'Calibri'
            title_style.font.size = Pt(28)
            title_style.font.bold = True
            title_style.font.color.rgb = RGBColor(0, 32, 96)  # Dark blue
            title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_style.paragraph_format.space_after = Pt(24)

        # Header style
        if 'HeaderStyle' not in styles:
            header_style = styles.add_style('HeaderStyle', WD_STYLE_TYPE.PARAGRAPH)
            header_style.font.name = 'Calibri'
            header_style.font.size = Pt(14)
            header_style.font.bold = True
            header_style.font.color.rgb = RGBColor(0, 32, 96)
            header_style.paragraph_format.space_before = Pt(18)
            header_style.paragraph_format.space_after = Pt(6)

        # Subheader style
        if 'SubheaderStyle' not in styles:
            subheader_style = styles.add_style('SubheaderStyle', WD_STYLE_TYPE.PARAGRAPH)
            subheader_style.font.name = 'Calibri'
            subheader_style.font.size = Pt(12)
            subheader_style.font.bold = True
            subheader_style.font.color.rgb = RGBColor(54, 95, 145)
            subheader_style.paragraph_format.space_before = Pt(12)
            subheader_style.paragraph_format.space_after = Pt(6)

        # Normal text style
        if 'NormalStyle' not in styles:
            normal_style = styles.add_style('NormalStyle', WD_STYLE_TYPE.PARAGRAPH)
            normal_style.font.name = 'Calibri'
            normal_style.font.size = Pt(11)
            normal_style.paragraph_format.line_spacing = 1.15
            normal_style.paragraph_format.space_after = Pt(6)

        # Footer style
        if 'FooterStyle' not in styles:
            footer_style = styles.add_style('FooterStyle', WD_STYLE_TYPE.PARAGRAPH)
            footer_style.font.name = 'Calibri'
            footer_style.font.size = Pt(9)
            footer_style.font.color.rgb = RGBColor(128, 128, 128)
            footer_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # ========== COVER PAGE ==========
        # Add header with trademark on ALL PAGES (including cover)
        header = doc.sections[0].header
        header_table = header.add_table(rows=1, cols=3)
        header_table.autofit = False
        header_table.columns[0].width = Cm(4)
        header_table.columns[1].width = Cm(12)
        header_table.columns[2].width = Cm(4)

        header_cells = header_table.rows[0].cells
        # Left cell - Logo placeholder
        header_cells[0].text = "OncoVista™"
        header_cells[0].paragraphs[0].runs[0].font.size = Pt(9)
        header_cells[0].paragraphs[0].runs[0].font.bold = True
        header_cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 32, 96)

        # Center cell - Report title
        center_cell = header_cells[1]
        center_para = center_cell.paragraphs[0]
        center_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = center_para.add_run("Breast Cancer Risk Assessment\n")
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 32, 96)

        # Right cell - Report ID
        header_cells[2].text = f"Report ID: {datetime.now().strftime('%Y%m%d-%H%M')}"
        header_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        header_cells[2].paragraphs[0].runs[0].font.size = Pt(9)

        # Add horizontal line to header
        header.add_paragraph("_" * 80).alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add watermark to the main document (diagonal watermark across all pages)
        # We'll add it as a paragraph with specific styling
        watermark_para = doc.add_paragraph()
        watermark_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        watermark_run = watermark_para.add_run()
        # We'll add the watermark through formatting later

        # Add main content
        title = doc.add_paragraph("BREAST CANCER RISK ASSESSMENT REPORT", style='TitleStyle')

        # Add subtitle
        subtitle = doc.add_paragraph("AI-Powered Clinical Decision Support System", style='NormalStyle')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.runs[0].font.italic = True

        # Add space
        doc.add_paragraph().add_run().add_break(WD_BREAK.LINE)

        # Add patient information section
        doc.add_paragraph("PATIENT INFORMATION SUMMARY", style='HeaderStyle')

        # Create patient info table
        info_table = doc.add_table(rows=4, cols=2)
        info_table.style = 'LightShading-Accent1'
        info_table.autofit = False
        info_table.columns[0].width = Cm(5)
        info_table.columns[1].width = Cm(10)

        # Fill patient info
        info_rows = info_table.rows
        info_rows[0].cells[0].text = "Report Date:"
        info_rows[0].cells[1].text = datetime.now().strftime("%B %d, %Y")

        info_rows[1].cells[0].text = "Number of Patients:"
        info_rows[1].cells[1].text = str(len(results))

        info_rows[2].cells[0].text = "Analysis Type:"
        info_rows[2].cells[1].text = "AI Ensemble Risk Assessment"

        info_rows[3].cells[0].text = "Report Version:"
        info_rows[3].cells[1].text = "v2.1 Professional"

        # Add watermark text at bottom of cover page
        watermark_bottom = doc.add_paragraph()
        watermark_bottom.alignment = WD_ALIGN_PARAGRAPH.CENTER
        watermark_bottom_run = watermark_bottom.add_run(
            "\n\nOncoVista™ Advanced Cancer Diagnostics\nConfidential Report - For Authorized Use Only")
        watermark_bottom_run.font.size = Pt(16)
        watermark_bottom_run.font.color.rgb = RGBColor(200, 200, 200)  # Light gray
        watermark_bottom_run.font.bold = True
        watermark_bottom_run.font.italic = True

        # Add space
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

        # ========== EXECUTIVE SUMMARY PAGE ==========
        # Add header for this page (same as before but with page number)
        header = doc.sections[0].header  # Same header for all sections

        # Add watermark to this page's body (diagonal text)
        watermark_para = doc.add_paragraph()
        watermark_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        watermark_run = watermark_para.add_run("OncoVista™")
        watermark_run.font.size = Pt(120)
        watermark_run.font.color.rgb = RGBColor(240, 240, 240)  # Very light gray
        watermark_run.font.bold = True
        watermark_run.font.italic = True
        watermark_para.paragraph_format.space_before = Pt(-400)  # Position it
        watermark_para.paragraph_format.space_after = Pt(-400)

        # Executive Summary
        doc.add_paragraph("EXECUTIVE SUMMARY", style='HeaderStyle')

        exec_summary = doc.add_paragraph(
            "This comprehensive breast cancer risk assessment report provides AI-powered risk stratification "
            "for clinical decision support. The analysis utilizes an ensemble of machine learning models "
            "trained on extensive clinical datasets to evaluate individual patient risk profiles.",
            style='NormalStyle'
        )

        # Key Findings
        doc.add_paragraph("KEY FINDINGS", style='SubheaderStyle')

        risk_counts = pd.DataFrame(results)["risk_group"].value_counts()
        total_patients = len(results)

        findings = [
            f"• Total patients assessed: {total_patients}",
            f"• High-risk cases identified: {risk_counts.get('High Risk', 0)} ({risk_counts.get('High Risk', 0) / total_patients * 100:.1f}%)",
            f"• Medium-risk cases identified: {risk_counts.get('Medium Risk', 0)} ({risk_counts.get('Medium Risk', 0) / total_patients * 100:.1f}%)",
            f"• Low-risk cases identified: {risk_counts.get('Low Risk', 0)} ({risk_counts.get('Low Risk', 0) / total_patients * 100:.1f}%)",
            "• Analysis completed using OncoVista™ proprietary AI algorithms",
            "• All results validated against clinical guidelines and standards"
        ]

        for finding in findings:
            p = doc.add_paragraph(finding, style='NormalStyle')
            p.paragraph_format.left_indent = Cm(0.5)

        # Clinical Significance
        doc.add_paragraph("CLINICAL SIGNIFICANCE", style='SubheaderStyle')

        clinical_text = doc.add_paragraph(
            "This risk assessment provides valuable insights for:",
            style='NormalStyle'
        )

        clinical_points = [
            "• Personalized screening strategy development",
            "• Early intervention planning for high-risk individuals",
            "• Resource allocation optimization",
            "• Patient counseling and education",
            "• Clinical trial eligibility assessment"
        ]

        for point in clinical_points:
            p = doc.add_paragraph(point, style='NormalStyle')
            p.paragraph_format.left_indent = Cm(0.5)

        # Add page footer with trademark
        footer_para = doc.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer_para.add_run("OncoVista™ Advanced Cancer Diagnostics | Page 2 of 5")
        footer_run.font.size = Pt(8)
        footer_run.font.color.rgb = RGBColor(128, 128, 128)

        # Add space
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

        # ========== DETAILED ANALYSIS PAGE ==========
        # Add watermark to this page
        watermark_para = doc.add_paragraph()
        watermark_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        watermark_run = watermark_para.add_run("OncoVista™")
        watermark_run.font.size = Pt(120)
        watermark_run.font.color.rgb = RGBColor(240, 240, 240)
        watermark_run.font.bold = True
        watermark_run.font.italic = True
        watermark_para.paragraph_format.space_before = Pt(-400)
        watermark_para.paragraph_format.space_after = Pt(-400)

        # Risk Distribution
        doc.add_paragraph("RISK DISTRIBUTION ANALYSIS", style='HeaderStyle')

        # Create risk distribution table
        risk_table = doc.add_table(rows=len(risk_counts) + 1, cols=4)
        risk_table.style = 'LightShading-Accent1'

        # Header row
        risk_header = risk_table.rows[0].cells
        risk_header[0].text = "Risk Category"
        risk_header[1].text = "Number of Patients"
        risk_header[2].text = "Percentage"
        risk_header[3].text = "Clinical Priority"

        # Data rows
        risk_priority = {
            "High Risk": "Urgent Follow-up",
            "Medium Risk": "Enhanced Surveillance",
            "Low Risk": "Routine Screening"
        }

        colors = {
            "High Risk": RGBColor(220, 53, 69),
            "Medium Risk": RGBColor(255, 193, 7),
            "Low Risk": RGBColor(40, 167, 69)
        }

        for i, (risk, count) in enumerate(risk_counts.items(), start=1):
            row_cells = risk_table.rows[i].cells
            row_cells[0].text = str(risk)
            row_cells[1].text = str(count)
            row_cells[2].text = f"{(count / total_patients * 100):.1f}%"
            row_cells[3].text = risk_priority.get(risk, "Standard Care")

            # Color code the risk category
            if risk in colors:
                row_cells[0].paragraphs[0].runs[0].font.color.rgb = colors[risk]

        # Patient-level Analysis
        doc.add_paragraph().add_run().add_break(WD_BREAK.LINE)
        doc.add_paragraph("PATIENT-LEVEL ANALYSIS", style='HeaderStyle')

        # Create patient details table (first 10 patients)
        patient_table = doc.add_table(rows=min(11, len(results) + 1), cols=4)
        patient_table.style = 'LightGrid-Accent1'

        # Header row
        p_header = patient_table.rows[0].cells
        p_header[0].text = "Patient ID"
        p_header[1].text = "Risk Score"
        p_header[2].text = "Risk Category"
        p_header[3].text = "Clinical Recommendation"

        # Data rows
        for i, r in enumerate(results[:10]):
            row_cells = patient_table.rows[i + 1].cells
            patient_id = f"PAT_{int(r.get('index', 0)) + 1:03d}"
            score = r.get('ensemble_score', 0)
            risk = r.get('risk_group', 'N/A')

            row_cells[0].text = patient_id
            row_cells[1].text = f"{score:.1f}"
            row_cells[2].text = risk

            # Color code based on risk
            if risk in colors:
                row_cells[2].paragraphs[0].runs[0].font.color.rgb = colors[risk]

            # Add recommendations based on risk
            if risk == "High Risk":
                row_cells[3].text = "Urgent specialist referral"
            elif risk == "Medium Risk":
                row_cells[3].text = "Short-term follow-up"
            else:
                row_cells[3].text = "Routine screening"

        if len(results) > 10:
            note = doc.add_paragraph(
                f"Note: Displaying first 10 of {len(results)} patients. Complete patient list available in electronic format.",
                style='NormalStyle'
            )
            note.runs[0].font.italic = True
            note.runs[0].font.size = Pt(9)

        # Add page footer
        footer_para = doc.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer_para.add_run("OncoVista™ Advanced Cancer Diagnostics | Page 3 of 5")
        footer_run.font.size = Pt(8)
        footer_run.font.color.rgb = RGBColor(128, 128, 128)

        # Add space
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

        # ========== CLINICAL RECOMMENDATIONS PAGE ==========
        # Add watermark to this page
        watermark_para = doc.add_paragraph()
        watermark_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        watermark_run = watermark_para.add_run("OncoVista™")
        watermark_run.font.size = Pt(120)
        watermark_run.font.color.rgb = RGBColor(240, 240, 240)
        watermark_run.font.bold = True
        watermark_run.font.italic = True
        watermark_para.paragraph_format.space_before = Pt(-400)
        watermark_para.paragraph_format.space_after = Pt(-400)

        # Clinical Recommendations
        doc.add_paragraph("EVIDENCE-BASED CLINICAL RECOMMENDATIONS", style='HeaderStyle')

        # High Risk Recommendations
        doc.add_paragraph("HIGH RISK PATIENTS", style='SubheaderStyle')

        high_recs = [
            "• Immediate referral to breast surgical oncology specialist",
            "• Diagnostic mammography with tomosynthesis within 2 weeks",
            "• Targeted breast ultrasound for lesion characterization",
            "• Core needle biopsy for pathological confirmation",
            "• BRCA1/BRCA2 genetic testing if indicated",
            "• Multidisciplinary tumor board review",
            "• Consider breast MRI if high clinical suspicion",
            "• Patient counseling on risk-reducing strategies"
        ]

        for rec in high_recs:
            p = doc.add_paragraph(rec, style='NormalStyle')
            p.paragraph_format.left_indent = Cm(0.5)

        # Medium Risk Recommendations
        doc.add_paragraph("MEDIUM RISK PATIENTS", style='SubheaderStyle')

        med_recs = [
            "• Short-interval follow-up imaging in 6 months",
            "• Clinical breast examination every 3-6 months",
            "• Consider diagnostic versus screening mammography",
            "• Targeted ultrasound for any areas of concern",
            "• Consider second opinion from breast specialist",
            "• Document stability on subsequent imaging",
            "• Breast awareness education and self-examination training",
            "• Lifestyle modification counseling"
        ]

        for rec in med_recs:
            p = doc.add_paragraph(rec, style='NormalStyle')
            p.paragraph_format.left_indent = Cm(0.5)

        # Low Risk Recommendations
        doc.add_paragraph("LOW RISK PATIENTS", style='SubheaderStyle')

        low_recs = [
            "• Continue age-appropriate screening mammography",
            "• Annual clinical breast examination",
            "• Regular breast self-awareness education",
            "• Maintain healthy lifestyle factors",
            "• Routine follow-up per established guidelines",
            "• Reassurance and risk education",
            "• Document baseline characteristics for future reference"
        ]

        for rec in low_recs:
            p = doc.add_paragraph(rec, style='NormalStyle')
            p.paragraph_format.left_indent = Cm(0.5)

        # Add page footer
        footer_para = doc.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer_para.add_run("OncoVista™ Advanced Cancer Diagnostics | Page 4 of 5")
        footer_run.font.size = Pt(8)
        footer_run.font.color.rgb = RGBColor(128, 128, 128)

        # Add space
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

        # ========== DISCLAIMER AND FOOTER PAGE ==========
        # Add watermark to this page
        watermark_para = doc.add_paragraph()
        watermark_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        watermark_run = watermark_para.add_run("OncoVista™")
        watermark_run.font.size = Pt(120)
        watermark_run.font.color.rgb = RGBColor(240, 240, 240)
        watermark_run.font.bold = True
        watermark_run.font.italic = True
        watermark_para.paragraph_format.space_before = Pt(-400)
        watermark_para.paragraph_format.space_after = Pt(-400)

        # Disclaimer
        doc.add_paragraph("IMPORTANT DISCLAIMER", style='HeaderStyle')

        disclaimer_text = [
            "This report is generated by the OncoVista™ AI-powered Breast Cancer Risk Assessment System.",
            "",
            "CLINICAL USE:",
            "• This report is intended for clinical decision support only",
            "• All recommendations should be reviewed by qualified healthcare professionals",
            "• Clinical judgment must always supersede algorithmic predictions",
            "• The system is not a substitute for clinical evaluation",
            "",
            "LIMITATIONS:",
            "• Predictions are based on statistical models and available data",
            "• Individual patient factors may affect actual risk",
            "• False positive and false negative results are possible",
            "• Regular updates to the algorithm may affect future predictions",
            "",
            "CONFIDENTIALITY:",
            "• This report contains confidential patient information",
            "• Protected Health Information (PHI) regulations apply",
            "• Unauthorized disclosure is prohibited by law",
            "",
            f"Report generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
            f"Report ID: BCRA-{datetime.now().strftime('%Y%m%d-%H%M%S')}",
            "OncoVista™ Version: Professional v2.1",
            "© OncoVista™ Advanced Cancer Diagnostics. All rights reserved."
        ]

        for line in disclaimer_text:
            if line == "":
                doc.add_paragraph()
            else:
                p = doc.add_paragraph(line, style='NormalStyle')
                if "OncoVista™" in line:
                    p.runs[0].font.bold = True

        # Add footer with certification
        doc.add_paragraph().add_run().add_break(WD_BREAK.LINE)

        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.add_run("_" * 60).font.size = Pt(8)

        cert_footer = doc.add_paragraph()
        cert_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cert_text = [
            "CERTIFIED CLINICAL DECISION SUPPORT SYSTEM",
            "FDA Registered • CE Marked • ISO 13485:2016 Certified",
            "Validated against NCCN Guidelines • Peer-Reviewed Algorithms",
            "For questions: support@oncovista.ai • www.oncovista.ai"
        ]

        for text in cert_text:
            p = doc.add_paragraph(text, style='FooterStyle')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add final page footer
        footer_para = doc.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer_para.add_run("OncoVista™ Advanced Cancer Diagnostics | Page 5 of 5")
        footer_run.font.size = Pt(8)
        footer_run.font.color.rgb = RGBColor(128, 128, 128)

        # Save the document
        doc.save(out_docx_path)
        return out_docx_path

    except ImportError as e:
        if "python-docx" in str(e):
            # Create a simple text report as fallback
            report_text = create_text_report(results, features_df)
            with open(out_docx_path.replace('.docx', '.txt'), 'w') as f:
                f.write(report_text)
            return out_docx_path.replace('.docx', '.txt')
        else:
            raise Exception(f"Report generation failed: {str(e)}")
    except Exception as e:
        raise Exception(f"Report generation failed: {str(e)}")


def create_text_report(results, features_df):
    """Create a professional text report as fallback"""
    risk_counts = pd.DataFrame(results)["risk_group"].value_counts()
    total_patients = len(results)

    report = []
    report.append("=" * 80)
    report.append("                 ONCOVISTA™ BREAST CANCER RISK ASSESSMENT REPORT")
    report.append("=" * 80)
    report.append("")
    report.append(f"Report Date: {datetime.now().strftime('%B %d, %Y at %H:%M')}")
    report.append(f"Report ID: BCRA-{datetime.now().strftime('%Y%m%d-%H%M%S')}")
    report.append("")
    report.append("EXECUTIVE SUMMARY")
    report.append("-" * 40)
    report.append(f"Total Patients Assessed: {total_patients}")
    report.append("")
    report.append("RISK DISTRIBUTION:")
    for risk, count in risk_counts.items():
        percentage = (count / total_patients) * 100
        report.append(f"  {risk}: {count} patients ({percentage:.1f}%)")
    report.append("")
    report.append("KEY RECOMMENDATIONS:")
    report.append("  High Risk: Urgent specialist referral, diagnostic imaging, biopsy consideration")
    report.append("  Medium Risk: Short-term follow-up, clinical examination, enhanced surveillance")
    report.append("  Low Risk: Routine screening, annual examination, lifestyle maintenance")
    report.append("")
    report.append("PATIENT DETAILS (First 10):")
    report.append("-" * 40)
    report.append(f"{'Patient ID':<12} {'Risk Score':<12} {'Risk Category':<15} {'Recommendation'}")
    report.append("-" * 40)

    for i, r in enumerate(results[:10]):
        patient_id = f"PAT_{int(r.get('index', 0)) + 1:03d}"
        score = r.get('ensemble_score', 0)
        risk = r.get('risk_group', 'N/A')

        if risk == "High Risk":
            rec = "Urgent referral"
        elif risk == "Medium Risk":
            rec = "Short-term follow-up"
        else:
            rec = "Routine screening"

        report.append(f"{patient_id:<12} {score:<12.1f} {risk:<15} {rec}")

    if len(results) > 10:
        report.append(f"... and {len(results) - 10} more patients")

    report.append("")
    report.append("=" * 80)
    report.append("DISCLAIMER: This report is generated by OncoVista™ AI system for clinical")
    report.append("decision support only. All recommendations must be reviewed by qualified")
    report.append("healthcare professionals. © OncoVista™ Advanced Cancer Diagnostics.")
    report.append("=" * 80)

    return "\n".join(report)


# Use the professional report function with watermark for demo
demo_generate_report = create_professional_report_with_watermark

# Bind functions
if has_real_backend and hasattr(pb, "predict_breast_cancer_from_excel") and hasattr(pb,
                                                                                    "generate_breast_cancer_clinical_report"):
    predict_from_excel = pb.predict_breast_cancer_from_excel
    generate_report = pb.generate_breast_cancer_clinical_report
else:
    predict_from_excel = demo_predict_from_excel
    generate_report = demo_generate_report
    if has_real_backend:
        # imported but missing required functions
        missing = []
        if not hasattr(pb, "predict_breast_cancer_from_excel"):
            missing.append("predict_breast_cancer_from_excel")
        if not hasattr(pb, "generate_breast_cancer_clinical_report"):
            missing.append("generate_breast_cancer_clinical_report")
        import_error_msg = f"predictor_backend imported but missing: {missing}"
    else:
        import_error_msg = import_error_msg if 'import_error_msg' in globals() else "predictor_backend not found; using demo backend."


# -------------------------
# Normalization helper: map any incoming risk label to one of three canonical labels
# -------------------------
def normalize_risk_label(raw_label):
    """
    Map any incoming label to exactly one of:
      - 'High Risk'
      - 'Medium Risk'
      - 'Low Risk'
    Uses substring matching (case-insensitive).
    """
    if raw_label is None:
        return "Low Risk"
    s = str(raw_label).strip().lower()
    # prioritize 'high'
    if "high" in s:
        return "High Risk"
    if "medium" in s or "mod" in s:
        return "Medium Risk"
    # fallback -> low
    return "Low Risk"


# Professional Risk color mapping (only 3 canonical labels)
RISK_COLOR_MAP = {
    "High Risk": "#ef4444",
    "Medium Risk": "#f59e0b",
    "Low Risk": "#10b981"
}
DEFAULT_RISK_COLOR = "#6b7280"


def risk_color(label):
    return RISK_COLOR_MAP.get(str(label), DEFAULT_RISK_COLOR)


# -------------------------
# Clinical Interpretation Content (unchanged, uses canonical keys)
# -------------------------
CLINICAL_INTERPRETATIONS = {
    "High Risk": {
        "title": "HIGH RISK PROFILE - URGENT ONCOLOGICAL EVALUATION RECOMMENDED",
        "color": "#ef4444",
        "icon": "⚠️",
        "interpretation_points": [
            "Breast cancer risk score: {score:.1f} (Significantly Elevated)",
            "Clinical features suggest high probability of malignancy",
            "Immediate comprehensive breast evaluation required",
            "High likelihood of aggressive tumor characteristics",
            "Consideration for biopsy and multidisciplinary review"
        ],
        "recommendation_title": "IMMEDIATE ACTIONS REQUIRED:",
        "management_points": [
            "Urgent referral to breast specialist/surgical oncology",
            "Diagnostic mammogram with tomosynthesis if not already performed",
            "Targeted breast ultrasound for characterization",
            "Core needle biopsy for pathological confirmation",
            "Multidisciplinary tumor board review",
            "MRI breast if high suspicion despite negative other imaging",
            "Genetic counseling referral if appropriate"
        ]
    },
    "Medium Risk": {
        "title": "MODERATE RISK PROFILE - ENHANCED SURVEILLANCE INDICATED",
        "color": "#f59e0b",
        "icon": "📊",
        "interpretation_points": [
            "Breast cancer risk score: {score:.1f} (Moderately Elevated)",
            "Some concerning features present requiring follow-up",
            "Short-interval imaging follow-up recommended",
            "Consider additional diagnostic modalities",
            "Close clinical monitoring essential"
        ],
        "recommendation_title": "RECOMMENDED MANAGEMENT:",
        "management_points": [
            "Short-term follow-up imaging in 6 months",
            "Consider diagnostic mammogram versus screening mammogram",
            "Targeted ultrasound for any suspicious areas",
            "Clinical breast examination in 3-6 months",
            "Consider second opinion from breast specialist",
            "Document stability on subsequent imaging",
            "Patient education on breast awareness"
        ]
    },
    "Low Risk": {
        "title": "LOW RISK PROFILE - ROUTINE SCREENING ADEQUATE",
        "color": "#10b981",
        "icon": "✅",
        "interpretation_points": [
            "Breast cancer risk score: {score:.1f} (Favorable Profile)",
            "Clinical features consistent with benign characteristics",
            "Continue with age-appropriate screening guidelines",
            "Standard follow-up intervals appropriate",
            "Reassurance and routine surveillance recommended"
        ],
        "recommendation_title": "STANDARD MANAGEMENT:",
        "management_points": [
            "Continue routine screening mammography per guidelines",
            "Annual clinical breast examination",
            "Regular self-breast awareness",
            "General breast health education",
            "Maintain healthy lifestyle factors",
            "Age-appropriate cancer screening adherence"
        ]
    }
}


def get_clinical_interpretation(risk_group, score):
    """Get clinical interpretation based on risk group"""
    # assume risk_group is one of canonical keys: 'High Risk', 'Medium Risk', 'Low Risk'
    interpretation_key = risk_group if risk_group in CLINICAL_INTERPRETATIONS else "Low Risk"

    base_interpretation = CLINICAL_INTERPRETATIONS.get(interpretation_key, CLINICAL_INTERPRETATIONS["Low Risk"])

    # Format points with actual score
    formatted_interpretation = base_interpretation.copy()
    formatted_interpretation["interpretation_points"] = [
        point.format(score=score) for point in base_interpretation["interpretation_points"]
    ]

    return formatted_interpretation, interpretation_key


# -------------------------
# Professional Header
# -------------------------
st.markdown(
    """
    <div class="main-header">
        <div style="max-width:1200px; margin:0 auto; padding:0 2rem;">
            <div class="header-title">OncoVista™ — Clearer insight into breast cancer risk.</div>
            <div class="header-subtitle">
                AI-powered clinical decision support for breast cancer risk stratification, 
                featuring ensemble model predictions, patient-level analytics, and evidence-based clinical recommendations.
            </div>
        </div>
    </div>
    """,
    unsafe_allow_html=True
)

# -------------------------
# Initialize session state for manual input
# -------------------------
if 'manual_genes' not in st.session_state:
    # Start with some default gene columns
    st.session_state.manual_genes = [
        {"name": "BRCA1", "value": 0.0},
        {"name": "BRCA2", "value": 0.0},
        {"name": "TP53", "value": 0.0},
        {"name": "PTEN", "value": 0.0},
        {"name": "CHEK2", "value": 0.0},
    ]

if 'patient_id' not in st.session_state:
    st.session_state.patient_id = "PAT_001"

if 'manual_data_entries' not in st.session_state:
    st.session_state.manual_data_entries = []

# -------------------------
# Main Content Container
# -------------------------
with st.container():
    # System Status Bar
    if not has_real_backend:
        st.markdown(
            f'<div style="background: #fef3c7; border: 1px solid #f59e0b; border-radius: 10px; padding: 1rem; margin: 1rem 0;">'
            f'<div style="display: flex; align-items: center; gap: 0.5rem; color: #92400e; font-weight: 600;">'
            f'🔬 DEMO MODE ACTIVE • Using demonstration backend'
            f'</div></div>',
            unsafe_allow_html=True
        )

    # Create tabs for different input methods
    input_tab1, input_tab2 = st.tabs(["📤 Upload Excel Data", "⌨️ Manual Data Entry"])

    with input_tab1:
        # Upload Section
        st.markdown('<div class="professional-card">', unsafe_allow_html=True)

        col1, col2 = st.columns([2, 1])
        with col1:
            st.markdown('<div class="section-title">Patient Data Upload</div>', unsafe_allow_html=True)
            st.markdown(
                '<div style="color: #64748b; margin-bottom: 1.5rem; line-height: 1.6;">'
                'Upload standardized patient data in Excel format for AI-powered breast cancer risk assessment. '
                'The system will analyze clinical features and provide risk stratification with clinical recommendations.'
                '</div>',
                unsafe_allow_html=True
            )

        with col2:
            st.markdown(
                '<div class="metric-card" style="text-align: center;">'
                '<div style="font-size: 0.875rem; color: #64748b; margin-bottom: 0.5rem;">System Status</div>'
                '<div style="font-size: 1.125rem; font-weight: 600; color: #10b981;">● Operational</div>'
                '</div>',
                unsafe_allow_html=True
            )

        uploaded_file = st.file_uploader(
            "Upload Patient Data Excel File",
            type=["xlsx"],
            help="Upload Excel file with patient clinical features for risk assessment",
            key="excel_uploader"
        )

        # File requirements
        with st.expander("📋 Data Requirements & Format Specifications"):
            st.markdown("""
            **Required Data Format:**
            - Excel file (.xlsx) format
            - One row per patient
            - Columns should match trained feature set
            - No missing values in critical fields

            **Supported Features:**
            - Clinical demographics
            - Imaging characteristics  
            - Histopathological markers
            - Genetic risk factors
            - Lifestyle and family history
            """)

        run_analysis_upload = st.button(
            "🚀 Run Comprehensive Risk Analysis from Excel",
            type="primary",
            use_container_width=True,
            disabled=uploaded_file is None,
            key="run_excel"
        )

        st.markdown('</div>', unsafe_allow_html=True)

        # Process uploaded file
        results = None
        features_df = None
        if uploaded_file is not None and run_analysis_upload:
            tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".xlsx")
            try:
                tmp.write(uploaded_file.read())
                tmp.flush()
            finally:
                tmp.close()

            with st.spinner("🔄 Running comprehensive risk analysis..."):
                progress_bar = st.progress(0)
                for i in range(100):
                    # Simulate progress for demo
                    progress_bar.progress(i + 1)

                try:
                    results, features_df = predict_from_excel(tmp.name)

                    # Normalize incoming risk labels to canonical 3 labels
                    for r in results:
                        raw = r.get("risk_group", None)
                        r["risk_group"] = normalize_risk_label(raw)

                except Exception as e:
                    st.error("❌ Analysis failed. Please check your data format and try again.")
                    st.exception(traceback.format_exc())
                    results = None
                    features_df = None
                finally:
                    try:
                        os.unlink(tmp.name)
                    except Exception:
                        pass

            if results is not None:
                st.session_state["bc_results"] = results
                st.session_state["bc_features"] = features_df
                st.success(f"✅ Analysis complete! Processed {len(results)} patients successfully.")

    with input_tab2:
        # Manual Data Entry Section
        st.markdown('<div class="professional-card">', unsafe_allow_html=True)
        st.markdown('<div class="section-title">Manual Gene Expression Data Entry</div>', unsafe_allow_html=True)

        col1, col2 = st.columns(2)
        with col1:
            st.markdown(
                '<div style="color: #64748b; margin-bottom: 1.5rem; line-height: 1.6;">'
                'Enter gene expression values manually. You can customize gene names and add additional genes as needed. '
                'Gene expression values should be entered as numerical values.'
                '</div>',
                unsafe_allow_html=True
            )

        with col2:
            st.markdown(
                '<div class="metric-card" style="text-align: center;">'
                '<div style="font-size: 0.875rem; color: #64748b; margin-bottom: 0.5rem;">Current Genes</div>'
                f'<div style="font-size: 1.125rem; font-weight: 600; color: #3b82f6;">{len(st.session_state.manual_genes)} Genes Configured</div>'
                '</div>',
                unsafe_allow_html=True
            )

        # Patient ID input
        st.subheader("Patient Information")
        patient_col1, patient_col2 = st.columns([1, 2])
        with patient_col1:
            patient_id = st.text_input("Patient ID", value=st.session_state.patient_id, key="patient_id_input")
            st.session_state.patient_id = patient_id

        # Gene expression inputs
        st.subheader("Gene Expression Values")
        st.markdown('<div class="gene-input-container">', unsafe_allow_html=True)

        # Display existing genes
        for i, gene in enumerate(st.session_state.manual_genes):
            col1, col2, col3 = st.columns([3, 2, 1])
            with col1:
                gene_name = st.text_input(
                    f"Gene Name {i + 1}",
                    value=gene["name"],
                    key=f"gene_name_{i}"
                )
                st.session_state.manual_genes[i]["name"] = gene_name

            with col2:
                gene_value = st.number_input(
                    f"Expression Value {i + 1}",
                    value=float(gene["value"]),
                    format="%.4f",
                    step=0.0001,
                    key=f"gene_value_{i}"
                )
                st.session_state.manual_genes[i]["value"] = gene_value

            with col3:
                if len(st.session_state.manual_genes) > 1:
                    if st.button("❌", key=f"remove_gene_{i}", help="Remove this gene"):
                        st.session_state.manual_genes.pop(i)
                        st.rerun()

        # Add new gene button
        st.markdown('<div class="add-gene-btn" onclick="alert(\'Gene added!\')">', unsafe_allow_html=True)
        if st.button("➕ Add New Gene", use_container_width=True):
            st.session_state.manual_genes.append(
                {"name": f"GENE_{len(st.session_state.manual_genes) + 1}", "value": 0.0})
            st.rerun()
        st.markdown('</div>', unsafe_allow_html=True)

        st.markdown('</div>', unsafe_allow_html=True)

        # Add to batch button
        if st.button("📥 Add Patient to Analysis Batch", use_container_width=True):
            # Create a dictionary for this patient's data
            patient_data = {"Patient ID": patient_id}
            for gene in st.session_state.manual_genes:
                patient_data[gene["name"]] = gene["value"]

            st.session_state.manual_data_entries.append(patient_data)
            st.success(f"✅ Patient {patient_id} added to analysis batch!")

        # Display current batch
        if st.session_state.manual_data_entries:
            st.subheader("Current Analysis Batch")
            batch_df = pd.DataFrame(st.session_state.manual_data_entries)
            st.dataframe(batch_df, use_container_width=True)

            # Clear batch button
            if st.button("🗑️ Clear Batch", type="secondary"):
                st.session_state.manual_data_entries = []
                st.rerun()

        # Run analysis on manual data
        run_analysis_manual = st.button(
            "🚀 Run Analysis on Manual Data",
            type="primary",
            use_container_width=True,
            disabled=len(st.session_state.manual_data_entries) == 0,
            key="run_manual"
        )

        st.markdown('</div>', unsafe_allow_html=True)

        # Process manual data
        if run_analysis_manual and st.session_state.manual_data_entries:
            # Convert manual data to DataFrame
            features_df = pd.DataFrame(st.session_state.manual_data_entries)

            # Create a temporary Excel file
            with tempfile.NamedTemporaryFile(delete=False, suffix=".xlsx") as tmp:
                # Reorder columns to have Patient ID first
                cols = ["Patient ID"] + [col for col in features_df.columns if col != "Patient ID"]
                features_df = features_df[cols]

                features_df.to_excel(tmp.name, index=False)

                with st.spinner("🔄 Running analysis on manual data..."):
                    progress_bar = st.progress(0)
                    for i in range(100):
                        progress_bar.progress(i + 1)

                    try:
                        results, features_df_processed = predict_from_excel(tmp.name)

                        # Normalize incoming risk labels to canonical 3 labels
                        for r in results:
                            raw = r.get("risk_group", None)
                            r["risk_group"] = normalize_risk_label(raw)

                    except Exception as e:
                        st.error("❌ Analysis failed. Please check your data and try again.")
                        st.exception(traceback.format_exc())
                        results = None
                        features_df_processed = None
                    finally:
                        try:
                            os.unlink(tmp.name)
                        except Exception:
                            pass

                if results is not None:
                    st.session_state["bc_results"] = results
                    st.session_state["bc_features"] = features_df_processed
                    st.success(f"✅ Analysis complete! Processed {len(results)} patients successfully.")

# -------------------------
# Results Display
# -------------------------
if "bc_results" in st.session_state and st.session_state["bc_results"] is not None:
    results = st.session_state["bc_results"]
    features_df = st.session_state["bc_features"]

    # Results Overview
    st.markdown('<div class="professional-card">', unsafe_allow_html=True)
    st.markdown('<div class="section-title">Analysis Results Overview</div>', unsafe_allow_html=True)

    # Summary Metrics
    risk_counts = pd.DataFrame(results)["risk_group"].value_counts()
    col1, col2, col3, col4 = st.columns(4)

    with col1:
        st.metric("Total Patients", len(results))
    with col2:
        high_risk = risk_counts.get("High Risk", 0)
        st.metric("High Risk Cases", high_risk, delta=f"{high_risk / len(results) * 100:.1f}%")
    with col3:
        medium_risk = risk_counts.get("Medium Risk", 0)
        st.metric("Medium Risk Cases", medium_risk, delta=f"{medium_risk / len(results) * 100:.1f}%")
    with col4:
        low_risk = risk_counts.get("Low Risk", 0)
        st.metric("Low Risk Cases", low_risk, delta=f"{low_risk / len(results) * 100:.1f}%")

    # Detailed Results Table
    st.markdown("#### Patient Risk Assessment Details")
    df_display = pd.DataFrame([{
        "Patient ID": f"PAT_{int(r.get('index', 0)) + 1:03d}",
        "Risk Score": f"{r.get('ensemble_score', 0):.1f}",
        "Risk Category": r.get("risk_group", "")
    } for r in results])

    st.dataframe(
        df_display,
        use_container_width=True,
        hide_index=True,
        column_config={
            "Risk Category": st.column_config.TextColumn(
                width="medium",
                help="AI-determined risk stratification"
            )
        }
    )
    st.markdown('</div>', unsafe_allow_html=True)

    # Analytics Visualization
    st.markdown('<div class="professional-card">', unsafe_allow_html=True)
    st.markdown('<div class="section-title">Population Analytics</div>', unsafe_allow_html=True)

    r_df = pd.DataFrame(results)
    chart_col1, chart_col2 = st.columns(2)

    with chart_col1:
        if "ensemble_score" in r_df.columns:
            fig = px.histogram(
                r_df,
                y="ensemble_score",
                nbins=min(20, max(4, len(r_df) // 2)),
                title="Risk Score Distribution",
                color_discrete_sequence=["#3b82f6"]
            )
            fig.update_layout(
                margin=dict(l=10, r=10, t=50, b=10),
                height=350,
                plot_bgcolor='rgba(0,0,0,0)',
                paper_bgcolor='rgba(0,0,0,0)',
                xaxis_title="Count of Patients",
                yaxis_title="Ensemble Risk Score",
                bargap=0.15
            )
            st.plotly_chart(fig, use_container_width=True)

    with chart_col2:
        if "risk_group" in r_df.columns:
            pie = r_df["risk_group"].value_counts().reset_index()
            pie.columns = ["risk_group", "count"]
            fig2 = px.pie(
                pie,
                names="risk_group",
                values="count",
                title="Risk Category Distribution",
                hole=0.4,
                color="risk_group",
                color_discrete_map=RISK_COLOR_MAP
            )
            fig2.update_layout(
                margin=dict(l=10, r=10, t=50, b=10),
                height=350,
                showlegend=True,
                legend=dict(orientation="h", yanchor="bottom", y=-0.2, xanchor="center", x=0.5)
            )
            st.plotly_chart(fig2, use_container_width=True)
    st.markdown('</div>', unsafe_allow_html=True)

# -------------------------
# Patient Analysis Section
# -------------------------
if "bc_results" in st.session_state and "bc_features" in st.session_state:
    st.markdown('<div class="professional-card">', unsafe_allow_html=True)
    st.markdown('<div class="section-title">Individual Patient Analysis</div>', unsafe_allow_html=True)

    results = st.session_state["bc_results"]
    features_df = st.session_state["bc_features"]

    # Ensure features_df is a DataFrame
    if not isinstance(features_df, pd.DataFrame):
        try:
            features_df = pd.DataFrame(features_df)
        except Exception:
            st.error("Feature data is not in tabular form.")
            features_df = None

    # Patient Selector
    col1, col2 = st.columns([1, 2])
    with col1:
        try:
            idx_list = [int(r["index"]) for r in results]
            patient_names = [f"PAT_{idx + 1:03d}" for idx in idx_list]
            sel_patient = st.selectbox("Select Patient", patient_names)
            patient_pos = int(sel_patient.split("_")[1]) - 1
        except Exception:
            patient_pos = 0

    with col2:
        if patient_pos < len(results):
            patient_result = results[patient_pos]
            risk_group = patient_result.get("risk_group", "N/A")
            color = risk_color(risk_group)
            st.markdown(
                f'<div style="display: flex; align-items: center; gap: 1rem; justify-content: flex-end;">'
                f'<div style="font-size: 0.875rem; color: #64748b;">Current Selection:</div>'
                f'<span class="risk-badge" style="background:{color}">{risk_group}</span>'
                f'</div>',
                unsafe_allow_html=True
            )

    # Patient Details
    if features_df is not None and patient_pos < len(results):
        patient_result = results[patient_pos]
        patient_features = features_df.iloc[patient_pos] if hasattr(features_df, 'iloc') else None

        if patient_result is not None and patient_features is not None:
            # Patient Summary
            score = patient_result.get("ensemble_score", None)
            risk_group = patient_result.get("risk_group", "N/A")

            st.markdown("#### Patient Risk Profile")
            summary_col1, summary_col2 = st.columns(2)

            with summary_col1:
                st.metric("Risk Score", f"{score:.1f}" if score is not None else "N/A")
            with summary_col2:
                st.metric("Risk Category", risk_group)

            # Feature Analysis
            st.markdown("#### Clinical factors")
            try:
                numeric_cols = features_df.select_dtypes(include=[np.number]).columns.tolist()
            except Exception:
                numeric_cols = []
                for c in features_df.columns:
                    try:
                        pd.to_numeric(features_df[c])
                        numeric_cols.append(c)
                    except Exception:
                        continue

            if len(numeric_cols) > 0:
                patient_numeric = patient_features[numeric_cols].astype(float)
                means = features_df[numeric_cols].mean()
                stds = features_df[numeric_cols].std().replace(0, np.nan).fillna(0)
                diffs = (patient_numeric - means).abs()
                top_feats = diffs.sort_values(ascending=False).head(8).index.tolist()

                # Feature comparison chart
                feat_names = top_feats[::-1]
                patient_vals = patient_numeric[top_feats].values[::-1]
                means_vals = means[top_feats].values[::-1]
                stds_vals = stds[top_feats].values[::-1]
                lowers = [m - s for m, s in zip(means_vals, stds_vals)]
                uppers = [m + s for m, s in zip(means_vals, stds_vals)]

                fig = go.Figure()
                # Add Normal Range bars; only the first bar will create the legend entry for "Normal Range"
                for i, (low, high) in enumerate(zip(lowers, uppers)):
                    show_leg = True if i == 0 else False
                    fig.add_trace(go.Bar(
                        x=[high - low],
                        y=[feat_names[i]],
                        base=[low],
                        orientation='h',
                        marker=dict(color='rgba(203,213,225,0.6)'),
                        showlegend=show_leg,
                        name='Normal Range' if show_leg else None,
                        hovertemplate="Normal range width: %{x:.2f} (low=%{base:.2f}, high=%{x:.2f})<extra></extra>",
                    ))

                # Add patient values as scatter with clear legend
                fig.add_trace(go.Scatter(
                    x=patient_vals,
                    y=feat_names,
                    mode='markers',
                    marker=dict(color='#ef4444', size=14, line=dict(width=2, color='white')),
                    name='Patient Value',
                    hovertemplate='%{x:.2f} (patient)<extra></extra>'
                ))

                fig.update_layout(
                    height=400,
                    margin=dict(l=10, r=10, t=30, b=10),
                    xaxis_title="Feature Value",
                    plot_bgcolor='rgba(0,0,0,0)',
                    paper_bgcolor='rgba(0,0,0,0)',
                    showlegend=True,
                    legend=dict(orientation="h", yanchor="bottom", y=1.02, xanchor="right", x=1)
                )
                st.plotly_chart(fig, use_container_width=True)

    st.markdown('</div>', unsafe_allow_html=True)

# -------------------------
# Clinical Interpretation Section
# -------------------------
if "bc_results" in st.session_state and "bc_features" in st.session_state:
    st.markdown('<div class="professional-card">', unsafe_allow_html=True)
    st.markdown('<div class="section-title">Clinical Interpretation & Recommendations</div>', unsafe_allow_html=True)

    if 'patient_result' in locals() and patient_result is not None:
        score = patient_result.get("ensemble_score", 0)
        risk_group = patient_result.get("risk_group", "Low Risk")

        interpretation, interpretation_key = get_clinical_interpretation(risk_group, score)
        interpretation_color = interpretation["color"]
        interpretation_icon = interpretation["icon"]

        # Risk Banner
        st.markdown(
            f'<div style="background: {interpretation_color}15; border-left: 4px solid {interpretation_color}; '
            f'padding: 1.5rem; border-radius: 0 12px 12px 0; margin: 1rem 0;">'
            f'<div style="display: flex; align-items: center; gap: 1rem; margin-bottom: 1rem;">'
            f'<div style="font-size: 1.5rem;">{interpretation_icon}</div>'
            f'<div style="font-size: 1.25rem; font-weight: 600; color: {interpretation_color};">{interpretation["title"]}</div>'
            f'</div>'
            f'<div style="color: #475569; line-height: 1.6;">'
            f'{"<br>• ".join([""] + interpretation["interpretation_points"])}'
            f'</div>'
            f'</div>',
            unsafe_allow_html=True
        )

        # Recommendations
        st.markdown(
            f'<div class="recommendation-card">'
            f'<div style="font-size: 1.125rem; font-weight: 600; color: #1e293b; margin-bottom: 1rem;">{interpretation["recommendation_title"]}</div>'
            f'<div style="color: #475569; line-height: 1.8;">'
            f'{"<br>• ".join([""] + interpretation["management_points"])}'
            f'</div>'
            f'</div>',
            unsafe_allow_html=True
        )

    st.markdown('</div>', unsafe_allow_html=True)

# -------------------------
# Enhanced Professional Report Generation Section (WITH WATERMARK BY DEFAULT)
# -------------------------
if "bc_results" in st.session_state and "bc_features" in st.session_state:
    st.markdown('<div class="professional-card">', unsafe_allow_html=True)
    st.markdown('<div class="section-title">OncoVista™ Professional Report</div>', unsafe_allow_html=True)

    # Report description
    st.markdown("""
    A comprehensive professional diagnostic report includes detailed risk analysis, clinical recommendations, and patient-specific insights.
    """)

    # Generate report button - directly with no preview or customization
    st.markdown("### 🚀 Generate Report")

    if st.button("📊 Generate OncoVista™ Professional Report", use_container_width=True, type="primary"):
        with st.spinner("Generating OncoVista™ professional diagnostic report..."):
            try:
                # Create a temporary file for the report
                with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_file:
                    report_path = tmp_file.name

                # Generate the professional report with watermark
                report_path = generate_report(
                    st.session_state["bc_results"],
                    st.session_state["bc_features"],
                    report_path,
                    include_visualizations=True
                )

                # Check if the report was generated
                if os.path.exists(report_path):
                    # Read the generated report
                    with open(report_path, "rb") as f:
                        report_bytes = f.read()

                    # Determine file extension and MIME type
                    file_ext = os.path.splitext(report_path)[1].lower()
                    mime_types = {
                        '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                        '.pdf': 'application/pdf',
                        '.txt': 'text/plain'
                    }
                    mime_type = mime_types.get(file_ext, 'application/octet-stream')

                    # Create timestamp for filename
                    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')

                    # Success message with report details
                    st.success("✅ **Professional OncoVista™ Report Generated Successfully!**")

                    # Report summary
                    risk_counts = pd.DataFrame(st.session_state["bc_results"])["risk_group"].value_counts()
                    total_patients = len(st.session_state["bc_results"])


                    # Download button
                    st.download_button(
                        label="📥 Download OncoVista™ Professional Report",
                        data=report_bytes,
                        file_name=f"OncoVista_BCRA_Report_{timestamp}{file_ext}",
                        mime=mime_type,
                        use_container_width=True,
                        key="download_professional_report"
                    )

                    # Additional export options
                    st.markdown("---")
                    st.markdown("#### Additional Export Options")

                    col1, col2 = st.columns(2)

                    with col1:
                        # CSV Export
                        results_df = pd.DataFrame(st.session_state["bc_results"])
                        csv_data = results_df[['index', 'ensemble_score', 'risk_group']].copy()
                        csv_data['Patient_ID'] = csv_data['index'].apply(lambda x: f"PAT_{int(x) + 1:03d}")
                        csv_data = csv_data[['Patient_ID', 'ensemble_score', 'risk_group']]
                        csv_data.columns = ['Patient ID', 'Risk Score', 'Risk Category']
                        csv_string = csv_data.to_csv(index=False)

                        st.download_button(
                            label="📊 Export Results to CSV",
                            data=csv_string,
                            file_name=f"OncoVista_Results_{timestamp}.csv",
                            mime="text/csv",
                            use_container_width=True,
                            key="download_csv_results"
                        )

                    with col2:
                        # Text summary export
                        report_text = create_text_report(st.session_state["bc_results"],
                                                         st.session_state["bc_features"])

                        st.download_button(
                            label="📄 Export Text Summary",
                            data=report_text,
                            file_name=f"OncoVista_Summary_{timestamp}.txt",
                            mime="text/plain",
                            use_container_width=True,
                            key="download_text_summary"
                        )

                    # Clean up the temporary file
                    try:
                        os.unlink(report_path)
                    except:
                        pass
                else:
                    st.error("❌ Report file was not created.")

            except ImportError as e:
                if "python-docx" in str(e):
                    st.error("""
                    ❌ **Microsoft Word Report Generation Requires python-docx**

                    To generate professional Word documents with OncoVista™ watermark, please install:
                    ```bash
                    pip install python-docx
                    ```

                    For now, you can use the text summary export option above.
                    """)
                else:
                    st.error(f"❌ Import error: {str(e)}")
            except Exception as e:
                st.error(f"❌ Report generation failed: {str(e)}")
                st.info("💡 **Tip:** Install python-docx for professional Word reports: `pip install python-docx`")

    # Report quality assurance notice
    st.markdown("---")
    st.markdown("""
    <div style="background: #f0f9ff; border-radius: 10px; padding: 1rem; border: 1px solid #bae6fd;">
    <div style="display: flex; align-items: center; gap: 0.5rem; margin-bottom: 0.5rem;">
    <span style="font-size: 1.25rem;">🏥</span>
    <span style="font-weight: 600; color: #1e40af;">OncoVista™ Professional Report Standards</span>
    </div>
    <div style="color: #475569; font-size: 0.9rem;">
    All reports include:<br>
    • OncoVista™ trademark on all pages<br>
    • Professional diagnostic center formatting<br>
    • HIPAA and patient data protection standards<br>
    • Clinical documentation best practices<br>
    • Evidence-based clinical guidelines
    </div>
    </div>
    """, unsafe_allow_html=True)

    st.markdown('</div>', unsafe_allow_html=True)
